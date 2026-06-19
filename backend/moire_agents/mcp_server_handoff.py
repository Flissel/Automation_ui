"""
MCP Server for Handoff Multi-Agent System

Exposes the handoff system to Claude Code for intelligent orchestration
of desktop automation tasks.

Usage:
    python mcp_server_handoff.py

Add to .claude/settings.json:
{
    "mcpServers": {
        "handoff": {
            "command": "python",
            "args": ["path/to/mcp_server_handoff.py"]
        }
    }
}

Production Deployment:
    See service/ directory for Windows service management and health monitoring.
"""

import asyncio
import json
import logging
import os
import signal
import sys
from datetime import datetime
from typing import Any, Dict, List, Optional

# Force-disable the JAX/Flax backends inside transformers BEFORE any
# transitive import pulls them in. On this machine jax has a broken
# _jax.dll that crashes SentenceTransformer().__init__() if transformers
# tries to probe for JAX. PyTorch is the only backend we actually need.
os.environ.setdefault("TRANSFORMERS_NO_JAX", "1")
os.environ.setdefault("USE_JAX", "0")
os.environ.setdefault("USE_FLAX", "0")
os.environ.setdefault("TRANSFORMERS_NO_ADVISORY_WARNINGS", "1")

# Add parent to path for imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Load backend/.env automatically so OPENROUTER_API_KEY and friends are
# available regardless of the working directory the MCP host starts us in.
# Without this, vision_analyze / PlanningTeam LLM calls fail silently with
# "No API key configured" even though the key is present in the repo.
try:
    from dotenv import load_dotenv as _load_dotenv  # type: ignore

    _env_candidates = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", ".env"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", ".env"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"),
    ]
    for _env_path in _env_candidates:
        _env_path = os.path.abspath(_env_path)
        if os.path.isfile(_env_path):
            _load_dotenv(_env_path, override=False)
            break
except ImportError:
    pass  # python-dotenv optional — env may already be set by the host

# Load production config
try:
    from config import get_config, load_config

    _config = load_config()
except ImportError:
    _config = None

# Setup logging
logging.basicConfig(
    level=logging.INFO if not _config else getattr(logging, _config.log_level),
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger("HandoffMCP")

# MCP imports
try:
    from mcp.server import Server
    from mcp.server.stdio import stdio_server
    from mcp.types import TextContent, Tool
except ImportError:
    print("MCP package not installed. Run: pip install mcp", file=sys.stderr)
    Server = None
    stdio_server = None
    Tool = None
    TextContent = None

# Handoff system imports
from agents.handoff import AgentRuntime, PlanningTeam, UserTask, ValidationTeam
# Claude CLI import
from agents.orchestrator import ClaudeCLIWrapper

# Global instances (initialized lazily)
_planning_team: Optional[PlanningTeam] = None
_validation_team: Optional[ValidationTeam] = None
_runtime: Optional[AgentRuntime] = None
_claude_cli: Optional[ClaudeCLIWrapper] = None

# Trusted root of the adaptive-skill library (SKILL.md tree). Referenced by the
# document/perception entrypoints (H._SKILL_LIB_ROOT) and by the skill_* handlers
# below. Keep this an absolute, hard-coded path so the trust boundary is explicit.
_SKILL_LIB_ROOT = r"C:\Users\User\Desktop\Vibemind_V1\vibemind-os\skills"


def get_claude_cli() -> ClaudeCLIWrapper:
    """Get or create the Claude CLI wrapper (singleton)."""
    global _claude_cli
    if _claude_cli is None:
        _claude_cli = ClaudeCLIWrapper()
    return _claude_cli


async def get_planning_team() -> PlanningTeam:
    """Get or create the planning team (singleton)."""
    global _planning_team
    if _planning_team is None:
        _planning_team = PlanningTeam(max_debate_rounds=2, use_llm=True)
        await _planning_team.start()
    return _planning_team


async def get_validation_team() -> ValidationTeam:
    """Get or create the validation team (singleton)."""
    global _validation_team
    if _validation_team is None:
        _validation_team = ValidationTeam(confidence_threshold=0.6)
        await _validation_team.start()
    return _validation_team


async def get_runtime() -> AgentRuntime:
    """Get or create the agent runtime (singleton)."""
    global _runtime
    if _runtime is None:
        _runtime = AgentRuntime()
    return _runtime


# Tool definitions
TOOLS = [
    Tool(
        name="handoff_plan",
        description="Create a desktop automation plan using LLM-powered Planner + Critic. "
        "Returns a plan with steps, approval status, issues, and confidence score.",
        inputSchema={
            "type": "object",
            "properties": {
                "goal": {
                    "type": "string",
                    "description": "What to accomplish (e.g., 'open notepad and type hello')",
                },
                "context": {
                    "type": "object",
                    "description": "Additional context (e.g., {message: 'hello', user_feedback: '...'})",
                },
            },
            "required": ["goal"],
        },
    ),
    Tool(
        name="handoff_execute",
        description="Execute a plan's automation steps. Each step can be: hotkey, sleep, write, press, click, find_and_click.",
        inputSchema={
            "type": "object",
            "properties": {
                "plan": {
                    "type": "array",
                    "description": "List of plan steps to execute",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {"type": "string"},
                            "description": {"type": "string"},
                        },
                    },
                },
                "speed_factor": {
                    "type": "number",
                    "default": 1.0,
                    "description": "Multiplier for sleep durations. 1.0 = real time (default). Use 0.5 for the legacy 2x speedup, 2.0 to slow down.",
                },
                "goal": {
                    "type": "string",
                    "description": "Optional goal description. When supplied, the execution is recorded in memory so future calls with the same goal can re-use the steps via memory_hint.",
                },
                "recovery": {
                    "type": "boolean",
                    "default": False,
                    "description": "Phase 2 opt-in. When true, run via runtime.execute_with_recovery() which retries failed steps via RecoveryAgent and replans on demand. Default false keeps the old linear behavior.",
                },
            },
            "required": ["plan"],
        },
    ),
    Tool(
        name="handoff_validate",
        description="Find and validate a UI element on screen. Returns location coordinates and confidence.",
        inputSchema={
            "type": "object",
            "properties": {
                "target": {
                    "type": "string",
                    "description": "Element to find (e.g., 'chat input field', 'send button')",
                },
                "expected_state": {
                    "type": "object",
                    "description": "Optional expected screen state for validation",
                },
            },
            "required": ["target"],
        },
    ),
    Tool(
        name="handoff_action",
        description="Execute a direct automation action (hotkey, type, press, click, sleep, scroll).",
        inputSchema={
            "type": "object",
            "properties": {
                "action_type": {
                    "type": "string",
                    "enum": ["hotkey", "type", "press", "click", "sleep", "scroll"],
                    "description": "Type of action to perform",
                },
                "params": {
                    "type": "object",
                    "description": "Action parameters (keys, text, key, x/y, seconds, direction/amount for scroll)",
                    "properties": {
                        "keys": {
                            "type": "string",
                            "description": "For hotkey: keys like 'ctrl+alt+space'",
                        },
                        "text": {
                            "type": "string",
                            "description": "For type: text to type",
                        },
                        "key": {
                            "type": "string",
                            "description": "For press: key name like 'enter'",
                        },
                        "x": {
                            "type": "integer",
                            "description": "For click/scroll: x coordinate",
                        },
                        "y": {
                            "type": "integer",
                            "description": "For click/scroll: y coordinate",
                        },
                        "seconds": {
                            "type": "number",
                            "description": "For sleep: duration",
                        },
                        "direction": {
                            "type": "string",
                            "enum": ["up", "down"],
                            "description": "For scroll: direction",
                        },
                        "amount": {
                            "type": "integer",
                            "description": "For scroll: number of scroll clicks (default 3)",
                        },
                    },
                },
            },
            "required": ["action_type", "params"],
        },
    ),
    Tool(
        name="handoff_status",
        description="Get the status of the handoff system including runtime stats.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_read_screen",
        description="Capture a monitor (or region) and run OCR on it. Multi-monitor aware via mss: monitor_id=0 is primary, 1 is secondary, etc. The full list of physical monitors is always returned in the `monitors` field. If OCR returns empty text, the screenshot is saved to `desktop_screenshots/` and the path is returned in `screenshot_path` — the base64 image is NOT embedded by default (use include_screenshot=True if you really need it).",
        inputSchema={
            "type": "object",
            "properties": {
                "region": {
                    "type": "object",
                    "description": "Optional region to capture {x, y, width, height} in GLOBAL screen coordinates. May straddle monitors. If omitted, captures the full selected monitor.",
                    "properties": {
                        "x": {"type": "integer"},
                        "y": {"type": "integer"},
                        "width": {"type": "integer"},
                        "height": {"type": "integer"},
                    },
                },
                "monitor_id": {
                    "type": "integer",
                    "description": "Physical monitor index (0 = primary). Ignored when `region` is supplied.",
                    "default": 0,
                },
                "include_screenshot": {
                    "type": "boolean",
                    "description": "If true, embed the full base64 PNG inline. Default false (uses screenshot_path instead to stay under MCP token limits).",
                    "default": False,
                },
                "with_vision": {
                    "type": "boolean",
                    "description": "If true, also run a vision-model (Gemini via OpenRouter) analysis of the same screenshot and return it under `vision`, including a `semantic_score` (via sentence-transformers) and a `token_overlap_score` vs the extracted text. Lets the caller validate layout + semantics alongside raw text.",
                    "default": False,
                },
                "vision_prompt": {
                    "type": "string",
                    "description": "Optional custom prompt for the vision analysis. Defaults to a generic 'describe this screen' prompt.",
                },
                "prefer_uia": {
                    "type": "boolean",
                    "description": "When true (default), try UI-Automation text extraction first — dramatically cleaner than pixel OCR for native apps, browsers, and Electron. Falls back to Tesseract if UIA yields <200 chars.",
                    "default": True,
                },
                "text_source": {
                    "type": "string",
                    "enum": ["auto", "uia", "tesseract"],
                    "description": "Force a specific text acquisition layer. 'auto' uses the layered pipeline (UIA → Tesseract → MoireServer).",
                    "default": "auto",
                },
            },
        },
    ),
    Tool(
        name="handoff_get_focus",
        description="Get the currently active/focused window. Returns window title, handle, and process ID. Use this to verify the correct window is focused before typing.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_scroll",
        description="Scroll the mouse wheel up or down. Can scroll at current position or at a specific location.",
        inputSchema={
            "type": "object",
            "properties": {
                "direction": {
                    "type": "string",
                    "enum": ["up", "down"],
                    "description": "Direction to scroll",
                },
                "amount": {
                    "type": "integer",
                    "description": "Number of scroll clicks (default: 3). Positive values scroll the content, negative not supported - use direction instead.",
                },
                "x": {
                    "type": "integer",
                    "description": "Optional x coordinate to scroll at. If not provided, scrolls at current mouse position.",
                },
                "y": {
                    "type": "integer",
                    "description": "Optional y coordinate to scroll at. If not provided, scrolls at current mouse position.",
                },
            },
            "required": ["direction"],
        },
    ),
    # Claude CLI Tools
    Tool(
        name="claude_cli_run",
        description="Run a prompt via Claude CLI. Can use skills and output as JSON.",
        inputSchema={
            "type": "object",
            "properties": {
                "prompt": {
                    "type": "string",
                    "description": "The prompt to send to Claude",
                },
                "skill": {
                    "type": "string",
                    "description": "Optional skill name to use",
                },
                "output_format": {
                    "type": "string",
                    "enum": ["text", "json"],
                    "default": "text",
                    "description": "Output format: text or json",
                },
            },
            "required": ["prompt"],
        },
    ),
    Tool(
        name="claude_cli_skill",
        description="Execute a Claude Skill with inputs.",
        inputSchema={
            "type": "object",
            "properties": {
                "skill_name": {
                    "type": "string",
                    "description": "Name of the skill (without .md extension)",
                },
                "inputs": {
                    "type": "object",
                    "description": "Input parameters for the skill",
                },
                "ui_context": {
                    "type": "object",
                    "description": "Optional UI context for desktop automation skills",
                },
            },
            "required": ["skill_name", "inputs"],
        },
    ),
    Tool(
        name="claude_cli_status",
        description="Check if Claude CLI is installed and available.",
        inputSchema={"type": "object", "properties": {}},
    ),
    # Vision Analysis Tool
    Tool(
        name="vision_analyze",
        description="Analyze screenshot with Gemini Vision AI. Returns UI analysis, element locations, and suggested automation actions. Uses cached stream frames when available.",
        inputSchema={
            "type": "object",
            "properties": {
                "prompt": {
                    "type": "string",
                    "description": "What to analyze (e.g., 'Find all buttons', 'What is the current state?', 'Locate the login form')",
                },
                "mode": {
                    "type": "string",
                    "enum": [
                        "element_detection",
                        "state_analysis",
                        "task_planning",
                        "custom",
                    ],
                    "description": "Analysis mode: element_detection (find UI elements), state_analysis (describe screen state), task_planning (suggest next actions), custom (free-form analysis)",
                },
                "json_output": {
                    "type": "boolean",
                    "default": True,
                    "description": "Return structured JSON response when possible",
                },
                "monitor_id": {
                    "type": "integer",
                    "default": 0,
                    "description": "Monitor index for multi-monitor setups (0 = primary)",
                },
            },
            "required": ["prompt"],
        },
    ),
    # Clawdbot Messaging Tools
    Tool(
        name="clawdbot_send_message",
        description="Send a message to a contact via messaging platform (WhatsApp, Telegram, Discord, Signal, etc.). "
        "Resolves contact names with fuzzy matching. The message is routed through the Clawdbot Gateway.",
        inputSchema={
            "type": "object",
            "properties": {
                "recipient": {
                    "type": "string",
                    "description": "Contact name, alias, or key (e.g., 'Peter', 'boss', 'mama'). Supports fuzzy matching.",
                },
                "message": {
                    "type": "string",
                    "description": "Message text to send. Supports {variable} placeholders.",
                },
                "platform": {
                    "type": "string",
                    "enum": [
                        "whatsapp",
                        "telegram",
                        "discord",
                        "signal",
                        "imessage",
                        "email",
                    ],
                    "description": "Messaging platform to use. If omitted, uses the first available platform for the contact.",
                },
            },
            "required": ["recipient", "message"],
        },
    ),
    Tool(
        name="clawdbot_get_contacts",
        description="List or search contacts in the registry. Supports fuzzy matching by name or alias. "
        "Returns contact info including available messaging platforms.",
        inputSchema={
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query for fuzzy matching (e.g., 'Peter', 'boss'). If empty, lists all contacts.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Maximum number of results to return (default: 10)",
                },
            },
        },
    ),
    Tool(
        name="clawdbot_get_status",
        description="Get the Clawdbot bridge status including active sessions, connected platforms, and capabilities.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="clawdbot_get_variables",
        description="Get predefined variables and message templates. Variables can be used as {variable_name} in messages. "
        "Templates are predefined message formats.",
        inputSchema={
            "type": "object",
            "properties": {
                "include_templates": {
                    "type": "boolean",
                    "default": True,
                    "description": "Whether to include message templates in the response",
                }
            },
        },
    ),
    # Clawdbot Browser Tools
    Tool(
        name="clawdbot_browser_open",
        description="Open a URL in the browser via Clawdbot. Use this for opening websites.",
        inputSchema={
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": "URL to open (e.g., 'https://google.com', 'github.com')",
                }
            },
            "required": ["url"],
        },
    ),
    Tool(
        name="clawdbot_browser_search",
        description="Search the web for a query via Clawdbot. Opens a Google search.",
        inputSchema={
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query (e.g., 'weather Berlin', 'Python docs')",
                }
            },
            "required": ["query"],
        },
    ),
    Tool(
        name="clawdbot_browser_read_page",
        description="Read the content of the currently open browser page using OCR.",
        inputSchema={"type": "object", "properties": {}},
    ),
    # Clawdbot Reporting Tool
    Tool(
        name="clawdbot_report_findings",
        description="Report/send findings or gathered information to a contact via messaging or as callback. "
        "Use after browser searches, page reads, or any operation where you want to communicate results. "
        "If no recipient specified, sends via Clawdbot callback channel.",
        inputSchema={
            "type": "object",
            "properties": {
                "findings": {
                    "type": "string",
                    "description": "The information/results to report (text summary)",
                },
                "recipient": {
                    "type": "string",
                    "description": "Optional: contact name to send findings to (e.g., 'Peter', 'boss'). If omitted, sends via callback.",
                },
                "platform": {
                    "type": "string",
                    "enum": ["whatsapp", "telegram", "discord", "signal", "email"],
                    "description": "Optional: messaging platform.",
                },
                "title": {
                    "type": "string",
                    "description": "Optional: short title/subject for the report",
                },
            },
            "required": ["findings"],
        },
    ),
    # ─── Phase 1.1: Event Queue Tools ────────────────────────────────────
    Tool(
        name="handoff_event_add",
        description="Queue a new event in the runtime EventQueue. Returns the event_id for later status / cancel calls.",
        inputSchema={
            "type": "object",
            "properties": {
                "event_type": {
                    "type": "string",
                    "description": "Logical type / goal of the event",
                },
                "payload": {
                    "type": "object",
                    "description": "Arbitrary structured payload",
                },
                "priority": {
                    "type": "integer",
                    "default": 0,
                    "description": "Higher value = scheduled earlier",
                },
            },
            "required": ["event_type"],
        },
    ),
    Tool(
        name="handoff_event_status",
        description="Look up an event's current status (pending / executing / completed / failed / cancelled) by id.",
        inputSchema={
            "type": "object",
            "properties": {
                "event_id": {"type": "string"},
            },
            "required": ["event_id"],
        },
    ),
    Tool(
        name="handoff_event_list",
        description="List events in the runtime queue, optionally filtered by status.",
        inputSchema={
            "type": "object",
            "properties": {
                "status_filter": {
                    "type": "string",
                    "description": "Filter by status: pending | planning | executing | completed | failed | cancelled",
                },
                "limit": {"type": "integer", "default": 50},
            },
        },
    ),
    Tool(
        name="handoff_event_cancel",
        description="Cancel a queued or in-progress event by id.",
        inputSchema={
            "type": "object",
            "properties": {"event_id": {"type": "string"}},
            "required": ["event_id"],
        },
    ),
    Tool(
        name="handoff_batch_execute",
        description="Bulk-add multiple events with a shared batch_id. Each entry: {event_type, payload?, priority?}.",
        inputSchema={
            "type": "object",
            "properties": {
                "events": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "event_type": {"type": "string"},
                            "payload": {"type": "object"},
                            "priority": {"type": "integer"},
                        },
                        "required": ["event_type"],
                    },
                },
            },
            "required": ["events"],
        },
    ),
    # ─── Phase 1.2: User Interaction Tools ───────────────────────────────
    Tool(
        name="handoff_notify",
        description="Send a fire-and-forget notification to the Vibemind UI (no answer expected).",
        inputSchema={
            "type": "object",
            "properties": {
                "message": {"type": "string"},
                "title": {"type": "string"},
                "level": {
                    "type": "string",
                    "enum": ["info", "warning", "error", "success"],
                    "default": "info",
                },
                "metadata": {"type": "object"},
            },
            "required": ["message"],
        },
    ),
    Tool(
        name="handoff_clarify",
        description="Ask the user a clarifying question. Returns immediately with a clarify_id; poll handoff_clarify_check for the answer.",
        inputSchema={
            "type": "object",
            "properties": {
                "question": {"type": "string"},
                "options": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional list of multiple-choice answers",
                },
                "timeout_seconds": {"type": "integer", "default": 300},
                "metadata": {"type": "object"},
            },
            "required": ["question"],
        },
    ),
    Tool(
        name="handoff_clarify_check",
        description="Look up a previously asked clarify by id. Returns status (pending | answered | timeout) and answer if available.",
        inputSchema={
            "type": "object",
            "properties": {"clarify_id": {"type": "string"}},
            "required": ["clarify_id"],
        },
    ),
    # ─── Phase 1.3: File / System Tools ──────────────────────────────────
    Tool(
        name="handoff_shell",
        description="Run a command from the shell allowlist. Allowlist = inspection / build / git tools only. NEVER uses a real shell, so |, &&, redirection are not supported.",
        inputSchema={
            "type": "object",
            "properties": {
                "command": {"type": "string"},
                "timeout": {"type": "integer", "default": 30},
                "cwd": {"type": "string"},
            },
            "required": ["command"],
        },
    ),
    Tool(
        name="handoff_file_search",
        description="Recursive glob from `root` (default cwd). Returns matching paths.",
        inputSchema={
            "type": "object",
            "properties": {
                "pattern": {
                    "type": "string",
                    "description": "glob pattern, e.g. '*.py'",
                },
                "root": {"type": "string"},
                "max_results": {"type": "integer", "default": 200},
            },
            "required": ["pattern"],
        },
    ),
    Tool(
        name="handoff_file_open",
        description="Open a file with the OS default application (Explorer / Finder / xdg-open).",
        inputSchema={
            "type": "object",
            "properties": {"path": {"type": "string"}},
            "required": ["path"],
        },
    ),
    Tool(
        name="handoff_dir_list",
        description="List a directory's contents with size and mtime.",
        inputSchema={
            "type": "object",
            "properties": {"path": {"type": "string"}},
        },
    ),
    Tool(
        name="handoff_file_read",
        description="Read up to max_bytes (default 1MiB) of a text file.",
        inputSchema={
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "max_bytes": {"type": "integer", "default": 1048576},
                "encoding": {"type": "string", "default": "utf-8"},
            },
            "required": ["path"],
        },
    ),
    Tool(
        name="handoff_file_write",
        description="Write a text file. Requires confirm=True. Refuses system directories. Max 5MiB.",
        inputSchema={
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "content": {"type": "string"},
                "mode": {"type": "string", "enum": ["w", "a"], "default": "w"},
                "encoding": {"type": "string", "default": "utf-8"},
                "confirm": {"type": "boolean", "default": False},
                "create_parents": {"type": "boolean", "default": False},
            },
            "required": ["path", "content"],
        },
    ),
    Tool(
        name="handoff_process_list",
        description="List running processes via psutil. Optional name filter.",
        inputSchema={
            "type": "object",
            "properties": {
                "name_filter": {"type": "string"},
                "limit": {"type": "integer", "default": 100},
            },
        },
    ),
    Tool(
        name="handoff_process_kill",
        description="Terminate a process by PID. Refuses system PIDs (< 100) and the MCP server itself.",
        inputSchema={
            "type": "object",
            "properties": {
                "pid": {"type": "integer"},
                "force": {"type": "boolean", "default": False},
            },
            "required": ["pid"],
        },
    ),
    Tool(
        name="handoff_system_info",
        description="Return cpu / memory / disk / boot info via psutil + platform.",
        inputSchema={"type": "object", "properties": {}},
    ),
    # ─── Phase 1.4: Smart Elements Tools ─────────────────────────────────
    Tool(
        name="handoff_find_element",
        description="Find a UI element by text / type / proximity. Consults pattern store cache first, then live ValidationTeam.",
        inputSchema={
            "type": "object",
            "properties": {
                "text": {"type": "string"},
                "element_type": {"type": "string"},
                "near_text": {"type": "string"},
            },
        },
    ),
    Tool(
        name="handoff_scroll_to",
        description="Scroll the screen until target element is found, optionally click it.",
        inputSchema={
            "type": "object",
            "properties": {
                "target": {"type": "string"},
                "element_type": {"type": "string"},
                "then_click": {"type": "boolean", "default": False},
                "direction": {
                    "type": "string",
                    "enum": ["up", "down"],
                    "default": "down",
                },
                "max_scrolls": {"type": "integer", "default": 10},
            },
            "required": ["target"],
        },
    ),
    # ─── Phase 1.5: Document Processing Tools ────────────────────────────
    Tool(
        name="handoff_doc_scan",
        description="Parse a PDF/DOCX/TXT into a structured page+section tree. Returns a document_id for later edit/apply/export calls.",
        inputSchema={
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "max_pages": {"type": "integer", "default": 20},
                "detect_structure": {"type": "boolean", "default": True},
            },
            "required": ["path"],
        },
    ),
    Tool(
        name="handoff_doc_edit",
        description="Queue an in-memory edit on a parsed document. Operations: replace | append | prepend | delete.",
        inputSchema={
            "type": "object",
            "properties": {
                "document_id": {"type": "string"},
                "page": {"type": "integer"},
                "section_index": {"type": "integer"},
                "new_text": {"type": "string"},
                "operation": {
                    "type": "string",
                    "enum": ["replace", "append", "prepend", "delete"],
                    "default": "replace",
                },
            },
            "required": ["document_id", "page", "section_index", "new_text"],
        },
    ),
    Tool(
        name="handoff_doc_apply",
        description="Write all queued edits back to the source file. dry_run=True returns the planned edit list without writing.",
        inputSchema={
            "type": "object",
            "properties": {
                "document_id": {"type": "string"},
                "dry_run": {"type": "boolean", "default": False},
            },
            "required": ["document_id"],
        },
    ),
    Tool(
        name="handoff_doc_export",
        description="Export the parsed document tree as JSON / Markdown / TXT.",
        inputSchema={
            "type": "object",
            "properties": {
                "document_id": {"type": "string"},
                "format": {
                    "type": "string",
                    "enum": ["json", "markdown", "txt"],
                    "default": "json",
                },
            },
            "required": ["document_id"],
        },
    ),
    Tool(
        name="handoff_doc_list",
        description="List all open document sessions with their pending-edit counts.",
        inputSchema={"type": "object", "properties": {}},
    ),
    # ─── Phase 4.1: eyeTerm Collaborative Tools ──────────────────────────
    Tool(
        name="handoff_eyeterm_status",
        description="Return eyeTerm's current state. Use to check availability before asking the user to look anywhere.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_eyeterm_get_gaze",
        description="Return the current gaze estimate (x, y, confidence). Returns available=false if eyeTerm /gaze is not implemented yet.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_eyeterm_dwell_request",
        description="Ask the user to look at a target for dwell_ms, return the dwelled point. Polls /gaze until a stationary cluster is detected.",
        inputSchema={
            "type": "object",
            "properties": {
                "prompt": {
                    "type": "string",
                    "description": "Question shown to the user via notify",
                },
                "dwell_ms": {"type": "integer", "default": 2000},
                "timeout_ms": {"type": "integer", "default": 15000},
                "confidence_threshold": {"type": "number", "default": 0.6},
            },
            "required": ["prompt"],
        },
    ),
    Tool(
        name="handoff_eyeterm_calibrate",
        description="Trigger eyeTerm's GA calibration loop. Returns available=false if /calibrate is not implemented yet.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_collaborative_select",
        description="High-level point-of-interest tool. Asks the user where, returns coordinates. mode='click' uses pynput, mode='gaze' uses eyeTerm with click fallback.",
        inputSchema={
            "type": "object",
            "properties": {
                "question": {"type": "string"},
                "mode": {
                    "type": "string",
                    "enum": ["click", "gaze"],
                    "default": "click",
                },
                "timeout_seconds": {"type": "integer", "default": 30},
            },
            "required": ["question"],
        },
    ),
    # ─── Phase 4.2: Structured Screen Description (for blind users) ──────
    Tool(
        name="handoff_describe_screen",
        description="Composite screen description: active window, all windows, physical monitors, regions, focused element, selection, cursor. Built so an agent can navigate without interpreting pixels. detail='full' adds OCR text for every monitor; detail='deep' additionally runs vision analysis per monitor and returns a cross_validation summary (OCR vs vision agreement_score).",
        inputSchema={
            "type": "object",
            "properties": {
                "detail": {
                    "type": "string",
                    "enum": ["summary", "full", "deep"],
                    "default": "summary",
                },
            },
        },
    ),
    Tool(
        name="handoff_describe_focus",
        description="Detailed description of the currently focused element via UI Automation: name, control_type, bounds, supported actions, parent chain.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_get_window_tree",
        description="List all visible windows with hwnd, pid, process name, bounds, z-order, is_minimized.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_get_selection",
        description="What the user has currently selected (best-effort via clipboard). Returns text, source app, source window.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_get_cursor_context",
        description="Cursor position plus the UI element under the cursor with bounds and supported actions.",
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="handoff_list_actionable",
        description="Walk the UI Automation tree of the active (or specified) window and list clickable / typeable elements.",
        inputSchema={
            "type": "object",
            "properties": {
                "window_hwnd": {
                    "type": "integer",
                    "description": "Optional window handle. Defaults to the foreground window.",
                },
                "max_elements": {"type": "integer", "default": 100},
            },
        },
    ),
    # ─── Phase 4.3: Screen-Change Subscription (Live Mode) ───────────────
    Tool(
        name="handoff_subscribe_screen_changes",
        description="Start a background task that polls describe_screen every interval_ms and publishes a 'screen_change' event when the active window or focus changes. Returns a subscription_id.",
        inputSchema={
            "type": "object",
            "properties": {
                "interval_ms": {"type": "integer", "default": 500},
                "subscription_id": {
                    "type": "string",
                    "description": "Optional caller-supplied id; auto-generated otherwise.",
                },
            },
        },
    ),
    Tool(
        name="handoff_unsubscribe_screen_changes",
        description="Stop a previously-started screen-change subscription.",
        inputSchema={
            "type": "object",
            "properties": {"subscription_id": {"type": "string"}},
            "required": ["subscription_id"],
        },
    ),
    Tool(
        name="handoff_get_screen_changes",
        description="Return the most recent in-memory events from a subscription (last 25 changes).",
        inputSchema={
            "type": "object",
            "properties": {"subscription_id": {"type": "string"}},
            "required": ["subscription_id"],
        },
    ),
    # ─── Phase 8: Self-Awareness Tool ──────────────────────────────────
    Tool(
        name="handoff_self_info",
        description="Returns info about this MCP server itself: tool count, uptime, loaded hub components, Python version. The agent's self-awareness tool.",
        inputSchema={"type": "object", "properties": {}},
    ),
    # Phase 7 VibeMind tools moved to vibemind MCP (vibemind_mcp.py)
    # ─── Phase 6: Direct Office COM Automation ──────────────────────────
    Tool(
        name="handoff_excel_fill",
        description="Fill Excel cells directly via COM — zero race conditions. Each cell: {ref:'A1', value:'text', formula:'=SUMME(...)', bold:true, italic:true, bg_color:'#4472C4', align:'center', border:true}. One call replaces 60+ keyboard actions.",
        inputSchema={
            "type": "object",
            "properties": {
                "cells": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "ref": {
                                "type": "string",
                                "description": "Cell reference, e.g. 'A1' or 'A1:E1'",
                            },
                            "value": {"description": "Text or number value"},
                            "formula": {
                                "type": "string",
                                "description": "Formula (use SUMME for German Excel)",
                            },
                            "bold": {"type": "boolean"},
                            "italic": {"type": "boolean"},
                            "font_size": {"type": "integer"},
                            "font_color": {
                                "type": "string",
                                "description": "Hex color, e.g. '#FF0000'",
                            },
                            "bg_color": {
                                "type": "string",
                                "description": "Hex background color",
                            },
                            "number_format": {"type": "string"},
                            "align": {
                                "type": "string",
                                "enum": ["left", "center", "right"],
                            },
                            "border": {"type": "boolean"},
                        },
                        "required": ["ref"],
                    },
                },
                "auto_fit": {"type": "boolean", "default": True},
                "sheet_name": {"type": "string"},
            },
            "required": ["cells"],
        },
    ),
    Tool(
        name="handoff_excel_read",
        description="Read cell values from Excel via COM. Returns the value(s) at the given range.",
        inputSchema={
            "type": "object",
            "properties": {
                "range_ref": {
                    "type": "string",
                    "description": "Cell or range, e.g. 'A1' or 'A1:E8'",
                },
                "sheet_name": {"type": "string"},
            },
            "required": ["range_ref"],
        },
    ),
    Tool(
        name="handoff_word_write",
        description="Write paragraphs to Word directly via COM. Each paragraph: {text:'...', style:'Heading 1', bold:true, italic:true, align:'center'}. Style names are language-dependent (German: 'Überschrift 1').",
        inputSchema={
            "type": "object",
            "properties": {
                "paragraphs": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "text": {"type": "string"},
                            "style": {"type": "string"},
                            "bold": {"type": "boolean"},
                            "italic": {"type": "boolean"},
                            "align": {
                                "type": "string",
                                "enum": ["left", "center", "right"],
                            },
                            "font_size": {"type": "integer"},
                        },
                        "required": ["text"],
                    },
                },
                "clear_first": {"type": "boolean", "default": False},
            },
            "required": ["paragraphs"],
        },
    ),
    Tool(
        name="handoff_office_save",
        description="Save the active Office document (Excel workbook or Word document).",
        inputSchema={
            "type": "object",
            "properties": {
                "app": {
                    "type": "string",
                    "enum": ["auto", "excel", "word"],
                    "default": "auto",
                },
            },
        },
    ),
    # ─── Phase 5: Geometry OCR (PaddleOCR) ───────────────────────────────
    Tool(
        name="handoff_read_screen_geometry",
        description="Extract text WITH precise polygon bounding-boxes from a monitor or region. Returns tokens[] and lines[] each with geometry (4-point polygon + bbox_xyxy in pixel coords), plus click_targets[] with the center (x,y) of each detected text region — directly usable for handoff_action(type='click'). Engine fallback: PaddleOCR → Tesseract+TSV → EasyOCR.",
        inputSchema={
            "type": "object",
            "properties": {
                "monitor_id": {
                    "type": "integer",
                    "default": 0,
                    "description": "Physical monitor index (0 = primary)",
                },
                "region": {
                    "type": "object",
                    "description": "Optional region {x, y, width, height} in global screen coordinates",
                    "properties": {
                        "x": {"type": "integer"},
                        "y": {"type": "integer"},
                        "width": {"type": "integer"},
                        "height": {"type": "integer"},
                    },
                },
                "engine": {
                    "type": "string",
                    "enum": ["auto", "rapidocr", "easyocr", "paddleocr", "tesseract"],
                    "default": "auto",
                    "description": "Force a specific OCR engine. 'auto' tries RapidOCR (PP-OCRv4 on ONNX) → EasyOCR → PaddleOCR → Tesseract.",
                },
                "return_char_boxes": {
                    "type": "boolean",
                    "default": False,
                    "description": "Request per-character boxes from PaddleOCR (very precise but verbose).",
                },
                "include_normalized": {
                    "type": "boolean",
                    "default": True,
                    "description": "Include geometry_normalized (0..1 coordinates) alongside pixel coords.",
                },
            },
        },
    ),
]


# Tool handlers
async def handle_plan(goal: str, context: Optional[Dict] = None) -> Dict[str, Any]:
    """Create a plan using PlanningTeam.

    Phase 0.4: before calling the LLM, look up similar successful tasks in
    persistent memory. If found, the previously-working steps are passed
    into the planner via context['memory_hint'] so the planner can prefer
    or adapt them rather than re-deriving from scratch.
    """
    context = dict(context or {})

    runtime = await get_runtime()
    cache_hit = False
    similar_count = 0
    if runtime.memory is not None:
        try:
            similar = runtime.memory.get_similar_tasks(goal, limit=3)
            similar_count = len(similar)
            successful_steps = runtime.memory.get_successful_steps_for_task(goal)
            if successful_steps:
                context["memory_hint"] = {
                    "source": "sqlite_memory",
                    "previous_successful_steps": successful_steps,
                    "note": "These steps worked for a previous identical task. Prefer/adapt them.",
                }
                cache_hit = True
        except Exception as e:
            logger.debug(f"handle_plan memory lookup failed: {e}")

    team = await get_planning_team()
    result = await team.create_plan(goal, context=context)

    # Annotate result with cache info (rückwärtskompatibel — nur additive Felder).
    if isinstance(result, dict):
        result.setdefault("memory", {})
        result["memory"]["cache_hit"] = cache_hit
        result["memory"]["similar_tasks"] = similar_count
    return result


# ─── Phase 3.3: Pattern learning helpers ─────────────────────────────────────


def _plan_steps_to_action_steps(plan: List[Dict[str, Any]]):
    """Convert plan-step dicts into ActionStep objects for PatternStore.

    Returns an empty list if the conversion module is unavailable.
    """
    try:
        from learning import action_step as _as  # type: ignore
    except Exception as e:
        logger.debug(f"action_step module not loaded: {e}")
        return []

    steps = []
    for step in plan:
        t = step.get("type", "")
        try:
            if t == "hotkey":
                steps.append(_as.hotkey(step.get("keys", "")))
            elif t in {"write", "type"}:
                steps.append(_as.type_text(step.get("text", step.get("content", ""))))
            elif t == "press":
                steps.append(_as.press(step.get("key", "enter")))
            elif t == "click":
                steps.append(_as.click(int(step.get("x", 0)), int(step.get("y", 0))))
            elif t == "scroll":
                steps.append(
                    _as.scroll(
                        step.get("direction", "down"), int(step.get("amount", 3))
                    )
                )
            elif t == "sleep":
                steps.append(
                    _as.wait(float(step.get("duration", step.get("seconds", 1))))
                )
            # find_and_click is not directly representable; skip.
        except Exception as e:
            logger.debug(f"step→action_step conversion failed for {t}: {e}")
    return steps


async def _learn_from_execution(
    runtime,
    goal: Optional[str],
    plan: List[Dict[str, Any]],
    success: bool,
    duration_ms: float,
) -> None:
    """Phase 3.3: feed an executed plan into PatternStore + MemoryCollector."""
    if not goal:
        return

    if runtime.pattern_store is not None:
        try:
            actions = _plan_steps_to_action_steps(plan)
            if actions:
                runtime.pattern_store.learn_pattern(
                    task=goal,
                    actions=actions,
                    success=success,
                    duration_ms=duration_ms,
                )
        except Exception as e:
            logger.debug(f"pattern_store.learn_pattern failed: {e}")

    if runtime.memory_collector is not None:
        try:
            mc = runtime.memory_collector
            if hasattr(mc, "start_episode"):
                mc.start_episode(app_context="mcp", goal=goal)
            if hasattr(mc, "end_episode"):
                mc.end_episode(success=success)
        except Exception as e:
            logger.debug(f"memory_collector record failed: {e}")


async def _execute_one_step(
    step: Dict[str, Any],
    speed_factor: float = 1.0,
) -> Dict[str, Any]:
    """Execute a single plan step. Extracted from handle_execute so the same
    code path can be used by Phase 2's execute_with_recovery loop.

    Returns a dict with at least {success: bool, type: str, error?: str}.
    Raises only for truly unexpected errors; expected step failures are
    captured in the returned dict.
    """
    import pyautogui
    import pyperclip

    step_type = step.get("type", "")
    step_result: Dict[str, Any] = {"type": step_type, "success": False}

    try:
        if step_type == "hotkey":
            keys = step.get("keys", "").replace("+", " ").split()
            if keys:
                pyautogui.hotkey(*keys)
            step_result["success"] = True

        elif step_type == "sleep":
            duration = step.get("duration", step.get("seconds", 1))
            actual = float(duration) * speed_factor
            await asyncio.sleep(actual)
            step_result["success"] = True

        elif step_type == "write":
            text = step.get("text", step.get("content", ""))
            if text:
                pyperclip.copy(text)
                pyautogui.hotkey("ctrl", "v")
            step_result["success"] = True

        elif step_type == "press":
            key = step.get("key", "enter")
            pyautogui.press(key)
            step_result["success"] = True

        elif step_type == "click":
            x = step.get("x", 0)
            y = step.get("y", 0)
            pyautogui.click(int(x), int(y))
            step_result["success"] = True

        elif step_type == "find_and_click":
            target = step.get("target", step.get("text", ""))
            if target:
                team = await get_validation_team()
                loc_result = await team.validate_element(target)
                if loc_result.get("element_location"):
                    loc = loc_result["element_location"]
                    pyautogui.click(loc["x"], loc["y"])
                    step_result["success"] = True
                    step_result["location"] = loc
                else:
                    step_result["error"] = f"Could not find '{target}'"
            else:
                step_result["error"] = "find_and_click without target"
        else:
            step_result["error"] = f"Unknown step type: {step_type}"

    except Exception as e:
        step_result["error"] = str(e)

    return step_result


async def handle_execute(
    plan: List[Dict[str, Any]],
    speed_factor: float = 1.0,
    goal: Optional[str] = None,
    recovery: bool = False,
) -> Dict[str, Any]:
    """Execute plan steps.

    Args:
        plan: list of step dicts
        speed_factor: multiplier for sleep durations. 1.0 means real time
            (default). Lower values speed execution up; higher values slow
            it down. Phase 0.3: replaces the previous hardcoded 0.5x global
            halving. Use speed_factor=0.5 explicitly to get the old behavior.
        goal: optional goal description. If provided, Phase 0.4 records the
            execution into AgentMemory so future plans for the same goal
            can short-circuit via memory_hint.
        recovery: Phase 2 opt-in. When True, run via
            runtime.execute_with_recovery() which retries failed steps via
            RecoveryAgent and replans via PlanningTeam.replan_from when
            asked. Default False keeps the old linear behavior.
    """
    if speed_factor <= 0:
        speed_factor = 1.0

    # Phase 0.4: open a memory task record if a goal was supplied.
    runtime = await get_runtime()
    memory_task_id: Optional[int] = None
    if goal and runtime.memory is not None:
        try:
            task_record = runtime.memory.start_task(goal, context={"steps": len(plan)})
            memory_task_id = task_record.id
        except Exception as e:
            # sqlite_memory.start_task hits a UNIQUE(task_hash) constraint when
            # the same goal has already been recorded. Fall back to attaching
            # this run's steps to the existing record instead of giving up.
            logger.info(
                f"handle_execute: start_task fell back to existing record ({e})"
            )
            try:
                similar = runtime.memory.get_similar_tasks(goal, limit=1)
                if similar:
                    memory_task_id = similar[0].id
            except Exception as inner:
                logger.warning(
                    f"handle_execute: similar-task fallback also failed ({inner})"
                )

    # Phase 2: opt-in recovery loop.
    _phase33_t0 = datetime.now().timestamp()
    if recovery:
        try:
            recovery_result = await runtime.execute_with_recovery(
                plan=plan,
                step_executor=lambda s: _execute_one_step(s, speed_factor),
                goal=goal,
            )
        except Exception as e:
            logger.warning(f"execute_with_recovery failed, falling back to linear: {e}")
            recovery = False
            recovery_result = None

        if recovery and recovery_result is not None:
            # Persist into memory if applicable.
            if memory_task_id is not None and runtime.memory is not None:
                try:
                    for step_record in recovery_result.get("results", []):
                        runtime.memory.add_task_step(
                            task_id=memory_task_id,
                            step_type=step_record.get("type", ""),
                            description=step_record.get("description", ""),
                            success=step_record.get("success", False),
                            details={
                                k: v for k, v in step_record.items() if k != "step"
                            },
                        )
                    runtime.memory.complete_task(
                        task_id=memory_task_id,
                        success=recovery_result.get("success", False),
                    )
                except Exception as e:
                    logger.debug(
                        f"handle_execute: memory persist after recovery failed: {e}"
                    )

            if memory_task_id is not None:
                recovery_result["memory_task_id"] = memory_task_id

            # Phase 3.3: feed into PatternStore + MemoryCollector.
            await _learn_from_execution(
                runtime=runtime,
                goal=goal,
                plan=plan,
                success=recovery_result.get("success", False),
                duration_ms=(datetime.now().timestamp() - _phase33_t0) * 1000,
            )
            return recovery_result

    results = []
    success = True

    for i, step in enumerate(plan):
        step_result = await _execute_one_step(step, speed_factor)
        step_result["step"] = i + 1
        if not step_result.get("success"):
            success = False

        # Small delay between steps (legacy behavior).
        await asyncio.sleep(0.15)

        results.append(step_result)

        # Phase 0.4: persist each step into memory.
        if memory_task_id is not None and runtime.memory is not None:
            try:
                runtime.memory.add_task_step(
                    task_id=memory_task_id,
                    step_type=step_type,
                    description=step.get("description", ""),
                    success=step_result.get("success", False),
                    details={k: v for k, v in step_result.items() if k != "step"},
                )
            except Exception as e:
                logger.debug(f"handle_execute memory.add_task_step failed: {e}")

    # Phase 0.4: close out the memory task record.
    if memory_task_id is not None and runtime.memory is not None:
        try:
            runtime.memory.complete_task(
                task_id=memory_task_id,
                success=success,
                error_message=None if success else "one or more steps failed",
            )
        except Exception as e:
            logger.debug(f"handle_execute memory.complete_task failed: {e}")

    response: Dict[str, Any] = {
        "success": success,
        "steps_executed": len(results),
        "results": results,
    }
    if memory_task_id is not None:
        response["memory_task_id"] = memory_task_id

    # Phase 3.3: feed into PatternStore + MemoryCollector.
    await _learn_from_execution(
        runtime=runtime,
        goal=goal,
        plan=plan,
        success=success,
        duration_ms=(datetime.now().timestamp() - _phase33_t0) * 1000,
    )
    return response


async def handle_validate(
    target: str, expected_state: Optional[Dict] = None
) -> Dict[str, Any]:
    """Validate a UI element."""
    team = await get_validation_team()
    result = await team.validate_element(target, expected_state)
    return result


async def handle_get_focus() -> Dict[str, Any]:
    """Get the currently active window."""
    from agents.handoff.window_focus import get_active_window

    return await get_active_window()


async def handle_set_focus(window_title: str) -> Dict[str, Any]:
    """Focus a window by (partial) title match."""
    from agents.handoff.window_focus import verify_window_focus

    return await verify_window_focus(window_title, timeout=3.0, auto_focus=True)


async def handle_list_windows() -> Dict[str, Any]:
    """List all visible windows with titles."""
    from agents.handoff.window_focus import list_visible_windows

    windows = list_visible_windows()
    return {"success": True, "windows": windows, "count": len(windows)}


async def handle_mouse_move(x: int, y: int, duration: float = 0.5) -> Dict[str, Any]:
    """Move mouse smoothly to position without clicking."""
    import pyautogui

    try:
        pyautogui.moveTo(int(x), int(y), duration=min(duration, 2.0))
        return {"success": True, "x": x, "y": y, "duration": duration}
    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_scroll(
    direction: str, amount: int = 3, x: Optional[int] = None, y: Optional[int] = None
) -> Dict[str, Any]:
    """Scroll the mouse wheel up or down."""
    import pyautogui

    try:
        # Move to position if specified
        if x is not None and y is not None:
            pyautogui.moveTo(x, y)

        # Determine scroll amount (positive = up, negative = down)
        clicks = abs(amount) if amount else 3
        if direction == "down":
            clicks = -clicks

        # Perform scroll
        pyautogui.scroll(clicks)

        return {
            "success": True,
            "action": "scroll",
            "direction": direction,
            "clicks": abs(amount) if amount else 3,
            "position": {"x": x, "y": y} if x is not None else "current",
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_action(action_type: str, params: Dict[str, Any]) -> Dict[str, Any]:
    """Execute a direct action with optional focus verification.

    Phase 0.3: `params['speed_factor']` controls the sleep duration multiplier
    (default 1.0 = real time). The previous hardcoded 0.5x global halving
    has been removed.
    """
    import pyautogui
    import pyperclip

    # Optional focus verification before action
    verify_focus = params.get("verify_focus", False)
    target_window = params.get("target_window", None)
    speed_factor = float(params.get("speed_factor", 1.0))
    if speed_factor <= 0:
        speed_factor = 1.0

    if verify_focus and target_window:
        from agents.handoff.window_focus import verify_window_focus

        focus_result = await verify_window_focus(target_window, timeout=3.0)
        if not focus_result["success"]:
            return {
                "success": False,
                "error": f"Window not focused: {target_window}",
                "focus_result": focus_result,
            }

    try:
        if action_type == "hotkey":
            keys = params.get("keys", "").replace("+", " ").split()
            if keys:
                pyautogui.hotkey(*keys)
            return {"success": True, "action": "hotkey", "keys": keys}

        elif action_type == "type":
            text = params.get("text", "")
            if text:
                pyperclip.copy(text)
                pyautogui.hotkey("ctrl", "v")
            return {"success": True, "action": "type", "text_length": len(text)}

        elif action_type == "press":
            key = params.get("key", "enter")
            pyautogui.press(key)
            return {"success": True, "action": "press", "key": key}

        elif action_type == "click":
            x = params.get("x", 0)
            y = params.get("y", 0)
            pyautogui.moveTo(int(x), int(y), duration=0.3)
            pyautogui.click()
            return {"success": True, "action": "click", "x": x, "y": y}

        elif action_type == "sleep":
            seconds = params.get("seconds", 1)
            actual = float(seconds) * speed_factor
            await asyncio.sleep(actual)
            return {
                "success": True,
                "action": "sleep",
                "seconds": seconds,
                "actual": actual,
                "speed_factor": speed_factor,
            }

        elif action_type == "scroll":
            direction = params.get("direction", "down")
            amount = params.get("amount", 3)
            x = params.get("x")
            y = params.get("y")

            # Move to position if specified
            if x is not None and y is not None:
                pyautogui.moveTo(x, y)

            # Scroll (positive = up, negative = down)
            clicks = abs(amount) if amount else 3
            if direction == "down":
                clicks = -clicks
            pyautogui.scroll(clicks)

            return {
                "success": True,
                "action": "scroll",
                "direction": direction,
                "clicks": abs(amount) if amount else 3,
                "position": {"x": x, "y": y} if x is not None else "current",
            }

        else:
            return {"success": False, "error": f"Unknown action type: {action_type}"}

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_status() -> Dict[str, Any]:
    """Get system status, including Phase 0 hub components and memory stats."""
    runtime = await get_runtime()
    stats = runtime.get_stats()

    response: Dict[str, Any] = {
        "runtime": {
            "tasks_processed": stats.get("tasks_processed", 0),
            "handoffs_routed": stats.get("handoffs_routed", 0),
            "errors": stats.get("errors", 0),
            "agents_registered": stats.get("agents_registered", 0),
            "active_sessions": stats.get("active_sessions", 0),
        },
        "planning_team": {
            "initialized": _planning_team is not None,
            "llm_enabled": _planning_team._use_llm if _planning_team else False,
        },
        "validation_team": {
            "initialized": _validation_team is not None,
            "confidence_threshold": (
                _validation_team.confidence_threshold if _validation_team else 0.6
            ),
        },
        # Phase 0.4: hub component status comes from runtime.get_stats()
        "hub": stats.get("hub", {}),
    }

    # Memory stats (only if memory was loaded at least once).
    if "memory_stats" in stats:
        response["memory_stats"] = stats["memory_stats"]

    # Event queue status (only if event_queue was loaded at least once).
    if "event_queue_status" in stats:
        response["event_queue_status"] = stats["event_queue_status"]

    return response


def _list_physical_monitors() -> List[Dict[str, int]]:
    """Enumerate real physical monitors via mss.

    mss.monitors[0] is the virtual combined display; [1:] are physical.
    We return the physical list indexed from 0 for API simplicity.
    """
    try:
        import mss  # type: ignore

        with mss.mss() as sct:
            return [
                {
                    "index": i,
                    "left": m["left"],
                    "top": m["top"],
                    "width": m["width"],
                    "height": m["height"],
                }
                for i, m in enumerate(sct.monitors[1:])
            ]
    except Exception as e:
        logger.debug(f"_list_physical_monitors failed: {e}")
        return []


#: Well-known install locations checked when tesseract is not on PATH.
_TESSERACT_STANDARD_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    r"C:\Users\%USERNAME%\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
    "/usr/bin/tesseract",
    "/usr/local/bin/tesseract",
    "/opt/homebrew/bin/tesseract",
]


def _tesseract_status() -> Dict[str, Any]:
    """Report whether Tesseract is usable in this environment.

    Resolution order: TESSERACT_PATH env var -> PATH via shutil.which ->
    well-known install locations. This fixes the common Windows case where
    the UB-Mannheim installer puts tesseract.exe in Program Files but does
    NOT add it to the user PATH.
    """
    import os as _os
    import shutil

    binary = _os.getenv("TESSERACT_PATH") or shutil.which("tesseract")
    if not binary:
        for candidate in _TESSERACT_STANDARD_PATHS:
            expanded = _os.path.expandvars(candidate)
            if _os.path.isfile(expanded):
                binary = expanded
                break
    try:
        import pytesseract  # type: ignore

        binding_ok = True
    except Exception:
        binding_ok = False
    return {
        "binding_installed": binding_ok,
        "binary_found": bool(binary),
        "binary_path": binary,
    }


#: Very common English / German words skipped from the agreement score so
#: noise like 'file', 'view', 'edit' in every OCR doesn't inflate scores.
_AGREEMENT_STOPWORDS = {
    "file",
    "edit",
    "view",
    "help",
    "tools",
    "window",
    "windows",
    "open",
    "close",
    "save",
    "name",
    "type",
    "text",
    "user",
    "users",
    "date",
    "time",
    "home",
    "back",
    "next",
    "prev",
    "menu",
    "find",
    "search",
    "click",
    "enter",
    "exit",
    "path",
    "with",
    "from",
    "this",
    "that",
    "these",
    "those",
    "when",
    "what",
    "where",
    "will",
    "have",
    "been",
    "your",
    "more",
    "less",
    "than",
    "also",
    "into",
    "over",
    "some",
    "such",
    "each",
    "item",
    "items",
    "list",
    "page",
    "tab",
    "tabs",
    "datei",
    "ansicht",
    "fenster",
    "bearbeiten",
    "öffnen",
    "speichern",
    "schließen",
    "suche",
    "name",
    "hilfe",
}


# ─── Embedding-based semantic score (Phase 4.4) ─────────────────────────────

# Singleton sentence-transformer model. Loaded lazily — first call pays the
# ~200ms load cost, every call after is ~15ms per embedding.
_ST_MODEL: Any = None
_ST_MODEL_LOAD_FAILED: bool = False


def _get_st_model():
    """Lazy-load the sentence-transformers model used for semantic scoring."""
    global _ST_MODEL, _ST_MODEL_LOAD_FAILED
    if _ST_MODEL is not None:
        return _ST_MODEL
    if _ST_MODEL_LOAD_FAILED:
        return None
    try:
        from sentence_transformers import SentenceTransformer  # type: ignore

        # all-MiniLM-L6-v2 is 80MB, 384-dim, fast, multilingual-ok. Already
        # used by la-fungus-search, so the weights are likely cached.
        _ST_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
        logger.info("Sentence-transformer model loaded: all-MiniLM-L6-v2")
        return _ST_MODEL
    except Exception as e:
        logger.warning(
            f"Sentence-transformer unavailable ({e}); falling back to token overlap"
        )
        _ST_MODEL_LOAD_FAILED = True
        return None


def _semantic_score(a: str, b: str) -> Optional[float]:
    """Cosine similarity between two texts in [0, 1]. None if model unavailable."""
    model = _get_st_model()
    if model is None or not a or not b:
        return None
    try:
        # Truncate extreme inputs — all-MiniLM caps at ~256 tokens anyway.
        a_in = a[:4000]
        b_in = b[:4000]
        emb = model.encode(
            [a_in, b_in], convert_to_numpy=True, normalize_embeddings=True
        )
        # Normalized embeddings -> cosine = dot product.
        sim = float((emb[0] * emb[1]).sum())
        # Cosine of normalized vectors is in [-1, 1]; clamp to [0, 1] for display.
        return round(max(0.0, min(1.0, (sim + 1) / 2)), 3)
    except Exception as e:
        logger.debug(f"_semantic_score failed: {e}")
        return None


def _agreement_score(ocr_text: str, vision_text: str) -> float:
    """Content-overlap score in [0, 1] tolerant to OCR noise.

    Strategy:
      1. Extract alpha-only tokens of length >= 5 from both sides
         (length 4 lets too much junk through; 5 is a sweet spot that keeps
         "chrome", "vscode", "explorer" but drops "file", "view").
      2. Drop common stopwords.
      3. Drop OCR tokens whose character class looks garbled
         (< 60% ASCII-letter ratio — Tesseract artifact detector).
      4. Compute what fraction of the remaining OCR tokens appear as a
         substring of the vision description (lowercased).

    Meant as a sanity check, not semantic similarity. Catches gross
    disagreements like "OCR saw login form" vs "vision saw image gallery".
    """
    if not ocr_text or not vision_text:
        return 0.0
    import re

    def _clean_tokens(text: str) -> set:
        raw = re.findall(r"[A-Za-zÀ-ÿ]{5,}", text)
        out = set()
        for tok in raw:
            lc = tok.lower()
            if lc in _AGREEMENT_STOPWORDS:
                continue
            # Drop tokens that are mostly non-ASCII latin garbage.
            ascii_ratio = sum(1 for c in tok if "a" <= c.lower() <= "z") / max(
                1, len(tok)
            )
            if ascii_ratio < 0.7:
                continue
            out.add(lc)
        return out

    ocr_tokens = _clean_tokens(ocr_text)
    if not ocr_tokens:
        return 0.0
    haystack = vision_text.lower()
    hits = sum(1 for t in ocr_tokens if t in haystack)
    return round(hits / len(ocr_tokens), 3)


async def handle_read_screen(
    region: Optional[Dict[str, int]] = None,
    monitor_id: int = 0,
    include_screenshot: bool = False,
    with_vision: bool = False,
    vision_prompt: Optional[str] = None,
    prefer_uia: bool = True,
    text_source: str = "auto",  # auto | uia | tesseract
) -> Dict[str, Any]:
    """Capture a monitor/region and extract text using a layered strategy.

    Text acquisition priority (when `text_source="auto"`):
      1. **UI Automation (UIA)** — structured text directly from every
         visible window that overlaps `monitor_id`. Dramatically cleaner
         than pixel OCR for native apps, browsers, and Electron apps.
         Only used when `prefer_uia=True` and no explicit `region` is set.
      2. **Tesseract OCR** — pixel fallback for canvas content, PDFs, or
         apps that don't expose a UIA tree.
      3. **MoireServer** — if :8766 is running, used as a last resort.

    Force a specific source with `text_source="uia"` or `"tesseract"`.

    Multi-monitor aware via `mss`. `monitor_id` maps to the N-th physical
    monitor (0 = primary). Pass `region={x,y,width,height}` in GLOBAL
    screen space for rectangle capture (UIA is then skipped).

    `include_screenshot=False` (default): only `screenshot_path` is
    returned — base64 inline would blow the MCP token budget.

    When `with_vision=True`, the same pixel buffer is sent to the vision
    agent (Gemini via OpenRouter) and two cross-validation scores are
    computed against the extracted text:
      - `token_overlap_score` — cheap stopword-filtered token match
      - `semantic_score` — cosine similarity via sentence-transformers
    The `semantic_score` is the authoritative one when available;
    `token_overlap_score` is kept for backwards compat.
    """
    import base64
    import io
    import time as _time

    from PIL import Image

    try:
        # ── 1. Acquire a PIL.Image for the requested area ───────────────
        screenshot = None
        from_cache = False

        # 1a. Try the live-streaming cache first.
        try:
            from stream_frame_cache import StreamFrameCache  # type: ignore

            cached_frame = StreamFrameCache.get_fresh_frame(
                monitor_id=monitor_id, max_age_ms=500
            )
            if cached_frame:
                screenshot = cached_frame.to_pil_image()
                from_cache = True
                logger.info(f"[read_screen] cached frame monitor={monitor_id}")
        except ImportError:
            pass
        except Exception as cache_error:
            logger.debug(f"[read_screen] cache lookup failed: {cache_error}")

        # 1b. Capture directly via mss if no cached frame.
        if screenshot is None:
            try:
                import mss  # type: ignore

                with mss.mss() as sct:
                    # mss monitors: [0] = virtual, [1:] = physical.
                    physical = sct.monitors[1:]
                    if region:
                        bbox = {
                            "left": int(region.get("x", 0)),
                            "top": int(region.get("y", 0)),
                            "width": int(region.get("width", 800)),
                            "height": int(region.get("height", 600)),
                        }
                    else:
                        if not physical:
                            return {
                                "success": False,
                                "error": "no monitors detected via mss",
                            }
                        target_idx = max(0, min(monitor_id, len(physical) - 1))
                        bbox = physical[target_idx]

                    raw = sct.grab(bbox)
                    screenshot = Image.frombytes(
                        "RGB", raw.size, raw.bgra, "raw", "BGRX"
                    )
            except ImportError:
                # Fall back to pyautogui (primary monitor only).
                import pyautogui  # type: ignore

                if region:
                    screenshot = pyautogui.screenshot(
                        region=(
                            int(region.get("x", 0)),
                            int(region.get("y", 0)),
                            int(region.get("width", 800)),
                            int(region.get("height", 600)),
                        )
                    )
                else:
                    screenshot = pyautogui.screenshot()

        # ── 2. Text acquisition (layered) ────────────────────────────────
        text_content = ""
        ocr_method = "none"
        tess = _tesseract_status()
        uia_summary: Optional[Dict[str, Any]] = None

        # Layer 1: UIA text extraction — skip when a raw region was requested
        # (UIA works on whole windows, not pixel rectangles).
        if prefer_uia and region is None and text_source in ("auto", "uia"):
            try:
                from agents.handoff.screen_description import \
                    extract_uia_text_for_monitor

                uia_result = await extract_uia_text_for_monitor(monitor_id=monitor_id)
                if uia_result.get("success"):
                    uia_text = uia_result.get("text", "").strip()
                    uia_summary = {
                        "window_count": uia_result.get("window_count", 0),
                        "total_chars": uia_result.get("total_chars", 0),
                        "per_window": [
                            {
                                "hwnd": w.get("hwnd"),
                                "title": w.get("title"),
                                "entry_count": w.get("entry_count"),
                                "char_count": w.get("char_count"),
                            }
                            for w in (uia_result.get("per_window") or [])
                        ],
                    }
                    # Heuristic: UIA is trusted when it yields >= 200 chars.
                    # Below that the window may be a canvas-only app (game,
                    # PDF viewer, YouTube player) where pixel OCR is better.
                    if len(uia_text) >= 200:
                        text_content = uia_text
                        ocr_method = "uia"
            except Exception as e:
                logger.debug(f"[read_screen] UIA layer failed: {e}")

        # Layer 2: Tesseract OCR (fallback or explicit request).
        if not text_content.strip() and text_source in ("auto", "tesseract"):
            if tess["binding_installed"] and tess["binary_found"]:
                try:
                    import pytesseract  # type: ignore

                    pytesseract.pytesseract.tesseract_cmd = tess["binary_path"]
                    text_content = pytesseract.image_to_string(screenshot) or ""
                    ocr_method = "pytesseract"
                except Exception as tess_error:
                    logger.debug(f"[read_screen] pytesseract error: {tess_error}")

        # Layer 3: MoireServer last-resort.
        if not text_content.strip() and text_source == "auto":
            try:
                from bridge.websocket_client import \
                    MoireWebSocketClient  # type: ignore

                client = MoireWebSocketClient(host="localhost", port=8766)
                await asyncio.wait_for(client.connect(), timeout=2.0)
                moire_result = await client.capture_and_wait_for_complete(timeout=10.0)
                if moire_result.success and moire_result.ui_context:
                    texts = [e.text for e in moire_result.ui_context.elements if e.text]
                    text_content = "\n".join(texts)
                    ocr_method = "moire_server"
                await client.disconnect()
            except Exception as moire_error:
                logger.debug(
                    f"[read_screen] MoireServer fallback failed: {moire_error}"
                )

        # ── 3. Build response (never embed base64 unless asked) ─────────
        result: Dict[str, Any] = {
            "success": True,
            "text": text_content,
            "text_length": len(text_content),
            "ocr_method": ocr_method,  # uia | pytesseract | moire_server | none
            "text_source": ocr_method,  # alias for clarity
            "from_cache": from_cache,
            "monitor_id": monitor_id,
            "screenshot_size": {
                "width": screenshot.width,
                "height": screenshot.height,
            },
            "tesseract": tess,
            "monitors": _list_physical_monitors(),
        }
        if uia_summary is not None:
            result["uia"] = uia_summary

        # If OCR produced nothing, save the screenshot to disk and point at it.
        if not text_content.strip():
            try:
                shots_dir = os.path.join(
                    os.path.dirname(os.path.abspath(__file__)), "desktop_screenshots"
                )
                os.makedirs(shots_dir, exist_ok=True)
                ts = _time.strftime("%Y%m%d_%H%M%S")
                path = os.path.join(shots_dir, f"screen_m{monitor_id}_{ts}.png")
                screenshot.save(path, format="PNG")
                result["screenshot_path"] = path
                if not tess["binary_found"]:
                    result["note"] = (
                        "Tesseract binary not found on PATH and no OCR backend "
                        "responded. Install Tesseract (Windows: "
                        "https://github.com/UB-Mannheim/tesseract/wiki) and set "
                        "TESSERACT_PATH, or run MoireServer on :8766."
                    )
                else:
                    result["note"] = (
                        "OCR returned empty text; screenshot saved to screenshot_path."
                    )
            except Exception as e:
                result["screenshot_save_error"] = str(e)

        if include_screenshot:
            try:
                buffer = io.BytesIO()
                screenshot.save(buffer, format="PNG")
                result["screenshot_base64"] = base64.b64encode(
                    buffer.getvalue()
                ).decode("utf-8")
            except Exception as e:
                result["screenshot_base64_error"] = str(e)

        # ── 4. Optional vision cross-check ─────────────────────────────
        if with_vision:
            vision_payload: Dict[str, Any] = {"available": False}

            # Error sentinels that vision_agent returns as plain strings
            # instead of raising. If the response starts with one of these,
            # treat the call as failed rather than counting gibberish.
            _VISION_ERROR_PREFIXES = (
                "Analyse-Fehler:",
                "Vision-Analyse nicht verfügbar",
                "Keine Vision-Backend verfügbar",
                "PIL nicht verfügbar",
                "vision error:",
            )

            try:
                from agents.vision_agent import \
                    get_vision_agent  # type: ignore

                agent = get_vision_agent()
                if agent is not None and agent.is_available():
                    # Serialize the screenshot once for the vision call.
                    buf = io.BytesIO()
                    screenshot.save(buf, format="PNG")
                    prompt = vision_prompt or (
                        "Describe this screen in detail. What application is "
                        "shown, what is the current state, what layout regions "
                        "exist (title bar, sidebar, main content, dialogs, "
                        "etc.), and what is the user most likely doing? "
                        "Keep the answer under 300 words."
                    )
                    # Retry once on transient provider errors (404/429/5xx
                    # from OpenRouter's upstream routing).
                    vision_text = ""
                    for attempt in range(2):
                        try:
                            vision_text = await agent.analyze_with_prompt(
                                buf.getvalue(), prompt
                            )
                        except Exception as e:
                            vision_text = f"vision error: {e}"

                        vt_stripped = (vision_text or "").strip()
                        transient = (
                            vt_stripped.startswith(
                                "Analyse-Fehler: OpenRouter API error: 4"
                            )
                            or vt_stripped.startswith(
                                "Analyse-Fehler: OpenRouter API error: 5"
                            )
                            or "Provider returned error" in vt_stripped
                        )
                        if not transient:
                            break
                        if attempt == 0:
                            await asyncio.sleep(0.8)  # brief backoff before retry

                    # Detect the error-string sentinels and fail honestly.
                    vt = (vision_text or "").strip()
                    if vt and any(vt.startswith(p) for p in _VISION_ERROR_PREFIXES):
                        vision_payload = {
                            "available": False,
                            "reason": vt,
                            "prompt": prompt,
                        }
                    else:
                        vision_payload = {
                            "available": True,
                            "description": vision_text,
                            "prompt": prompt,
                        }
                        if text_content and isinstance(vision_text, str):
                            # Both scores: token overlap (cheap, old) +
                            # semantic cosine (authoritative, new).
                            token_score = _agreement_score(text_content, vision_text)
                            sem_score = _semantic_score(text_content, vision_text)
                            vision_payload["token_overlap_score"] = token_score
                            if sem_score is not None:
                                vision_payload["semantic_score"] = sem_score
                                vision_payload["agreement_score"] = sem_score  # primary
                                vision_payload["score_method"] = "sentence_transformers"
                            else:
                                vision_payload["agreement_score"] = token_score
                                vision_payload["score_method"] = "token_overlap"
                else:
                    vision_payload["reason"] = (
                        "vision agent not initialised (missing OPENROUTER_API_KEY?)"
                    )
            except ImportError as e:
                vision_payload["reason"] = f"vision agent not importable: {e}"
            except Exception as e:
                vision_payload["reason"] = f"vision call failed: {e}"
            result["vision"] = vision_payload

        return result

    except Exception as e:
        return {"success": False, "error": str(e)}


# Claude CLI Handlers
async def handle_claude_run(
    prompt: str, skill: Optional[str] = None, output_format: str = "text"
) -> Dict[str, Any]:
    """Run a prompt via Claude CLI."""
    try:
        cli = get_claude_cli()

        if not cli.is_available():
            return {
                "success": False,
                "error": "Claude CLI not available. Install with: npm install -g @anthropic-ai/claude-cli",
            }

        # Build command args
        args = []
        if skill:
            args.extend(["--skill", skill])
        if output_format == "json":
            args.append("--output-format=json")

        # Run the command (async)
        result = await cli.run_command(prompt, skill=skill, output_format=output_format)

        return {
            "success": result.get("success", False),
            "output": result.get("output"),
            "error": result.get("error"),
            "prompt": prompt,
            "skill": skill,
            "output_format": output_format,
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_claude_skill(
    skill_name: str, inputs: Dict[str, Any], ui_context: Optional[Dict] = None
) -> Dict[str, Any]:
    """Execute a Claude Skill with inputs."""
    try:
        cli = get_claude_cli()

        if not cli.is_available():
            return {
                "success": False,
                "error": "Claude CLI not available. Install with: npm install -g @anthropic-ai/claude-cli",
            }

        # Run the skill (async)
        result = await cli.run_skill(skill_name, inputs, ui_context=ui_context)

        return {
            "success": result.get("success", False),
            "output": result.get("output"),
            "error": result.get("error"),
            "skill_name": skill_name,
            "inputs": inputs,
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_claude_status() -> Dict[str, Any]:
    """Check if Claude CLI is installed and available."""
    try:
        cli = get_claude_cli()

        available = cli.is_available()
        cli_path = cli.cli_path if hasattr(cli, "cli_path") else None

        result = {"success": True, "available": available, "cli_path": cli_path}

        if available:
            # Try to get version or additional info
            try:
                skills = cli.list_skills()
                result["skills_count"] = len(skills)
            except:
                pass

        return result

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_vision_analyze(
    prompt: str,
    mode: str = "custom",
    json_output: bool = True,
    monitor_id: int = 0,
    viewport: Dict[str, int] = None,
) -> Dict[str, Any]:
    """Analyze screenshot with Gemini Vision AI."""
    import base64
    import io

    try:
        # Try to import vision agent
        try:
            from agents.vision_agent import get_vision_agent

            vision_agent = get_vision_agent()
            has_vision = vision_agent is not None and vision_agent.is_available()
        except ImportError:
            has_vision = False
            vision_agent = None

        if not has_vision:
            return {
                "success": False,
                "error": "Vision agent not available. Check OpenRouter API key and vision_agent.py",
            }

        # Get screenshot from cache or capture directly
        screenshot_bytes = None
        frame_source = None

        # Try backend API first (cached frames from WebSocket stream)
        try:
            import httpx

            async with httpx.AsyncClient(timeout=2.0) as client:
                response = await client.get(
                    f"http://localhost:8007/api/desktop/cached-frame/{monitor_id}",
                    params={"max_age_ms": 500},
                )
                if response.status_code == 200:
                    data = response.json()
                    if data.get("success") and data.get("frame_data"):
                        frame_data = data["frame_data"]
                        # Remove data URL prefix if present
                        if frame_data.startswith("data:"):
                            frame_data = frame_data.split(",", 1)[1]
                        screenshot_bytes = base64.b64decode(frame_data)
                        frame_source = "api_cache"
                        logger.info(
                            f"Using cached frame from API for monitor {monitor_id} (age: {data.get('age_ms', 0):.0f}ms)"
                        )
        except Exception as e:
            logger.debug(f"Backend API cache not available: {e}")

        # Fallback: Try local StreamFrameCache (if MCP runs in same process)
        if screenshot_bytes is None:
            try:
                from stream_frame_cache import StreamFrameCache

                cached_frame = StreamFrameCache.get_fresh_frame(
                    monitor_id=monitor_id, max_age_ms=500
                )
                if cached_frame:
                    screenshot_bytes = cached_frame.to_bytes()
                    frame_source = "local_cache"
                    logger.info(f"Using local cached frame for monitor {monitor_id}")
            except Exception as e:
                logger.debug(f"Local StreamFrameCache not available: {e}")

        # Fallback to pyautogui screenshot
        if screenshot_bytes is None:
            import pyautogui

            screenshot = pyautogui.screenshot()
            buffer = io.BytesIO()
            screenshot.save(buffer, format="PNG")
            screenshot_bytes = buffer.getvalue()
            logger.info("Using pyautogui screenshot")

        # Crop to viewport if specified
        viewport_offset = None
        if viewport and screenshot_bytes:
            try:
                from PIL import Image

                img = Image.open(io.BytesIO(screenshot_bytes))
                vx = max(0, min(int(viewport.get("x", 0)), img.width - 1))
                vy = max(0, min(int(viewport.get("y", 0)), img.height - 1))
                vw = min(int(viewport.get("width", img.width)), img.width - vx)
                vh = min(int(viewport.get("height", img.height)), img.height - vy)
                if vw > 50 and vh > 50:  # Minimum 50x50
                    cropped = img.crop((vx, vy, vx + vw, vy + vh))
                    buffer = io.BytesIO()
                    cropped.save(buffer, format="PNG")
                    screenshot_bytes = buffer.getvalue()
                    viewport_offset = {"x": vx, "y": vy, "width": vw, "height": vh}
                    logger.info(
                        f"[vision_analyze] Cropped to viewport ({vx},{vy}) {vw}x{vh}"
                    )
            except Exception as e:
                logger.warning(f"[vision_analyze] Viewport crop failed: {e}")

        # Build analysis prompt based on mode
        mode_prompts = {
            "element_detection": f"Find all interactive UI elements (buttons, inputs, links, etc.) on this screen. {prompt}. Return as JSON with elements array containing: type, text, approximate_location (x, y), confidence.",
            "state_analysis": f"Analyze the current state of this screen. {prompt}. Describe what application is shown, what's visible, and the current state.",
            "task_planning": f"Based on this screen, suggest the next automation steps to accomplish: {prompt}. Return as JSON with steps array.",
            "custom": prompt,
        }

        analysis_prompt = mode_prompts.get(mode, prompt)

        # Inject viewport offset instructions so vision model reports absolute coordinates
        if viewport_offset:
            analysis_prompt += (
                f"\n\nIMPORTANT: This image is a CROPPED VIEWPORT from screen position "
                f"({viewport_offset['x']},{viewport_offset['y']}) size {viewport_offset['width']}x{viewport_offset['height']}. "
                f"When reporting element coordinates, ADD the offset: "
                f"absolute_x = element_x_in_image + {viewport_offset['x']}, "
                f"absolute_y = element_y_in_image + {viewport_offset['y']}. "
                f"All coordinates in your response MUST be absolute screen coordinates."
            )

        # Run vision analysis
        result = await vision_agent.analyze_with_prompt(
            screenshot_bytes, analysis_prompt
        )

        ret = {
            "success": True,
            "analysis": result,
            "mode": mode,
            "monitor_id": monitor_id,
            "source": frame_source if frame_source else "pyautogui",
        }
        if viewport_offset:
            ret["viewport"] = viewport_offset
        return ret

    except Exception as e:
        logger.error(f"Vision analysis failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "mode": mode,
            "monitor_id": monitor_id,
        }


# ============================================
# Clawdbot Tool Handlers
# ============================================

CLAWDBOT_API_BASE = "http://localhost:8007/api/clawdbot"


async def handle_clawdbot_send_message(
    recipient: str, message: str, platform: Optional[str] = None
) -> Dict[str, Any]:
    """Send a message to a contact via Clawdbot."""
    try:
        import httpx

        async with httpx.AsyncClient(timeout=10.0) as client:
            # Resolve contact first
            params = {}
            if platform:
                params["platform"] = platform

            resolve_resp = await client.post(
                f"{CLAWDBOT_API_BASE}/contacts/{recipient}/resolve", params=params
            )

            if resolve_resp.status_code == 404:
                return {
                    "success": False,
                    "error": f"Contact '{recipient}' not found",
                    "suggestion": "Use clawdbot_get_contacts to search for contacts",
                }

            resolve_data = resolve_resp.json()

            if not resolve_data.get("found"):
                return {
                    "success": False,
                    "error": f"Contact '{recipient}' not found",
                    "suggestions": resolve_data.get("suggestions", []),
                }

            contact = resolve_data["contact"]
            contact_name = contact.get("name", recipient)

            # Determine platform and recipient_id
            target_platform = platform
            recipient_id = None

            if platform:
                recipient_id = contact.get(platform.lower())

            if not recipient_id:
                for p in [
                    "whatsapp",
                    "telegram",
                    "discord",
                    "signal",
                    "imessage",
                    "email",
                ]:
                    if contact.get(p):
                        target_platform = p
                        recipient_id = contact[p]
                        break

            if not recipient_id:
                return {
                    "success": False,
                    "error": f"Contact '{contact_name}' has no messaging platform configured",
                }

            # Send via the command endpoint
            cmd_resp = await client.post(
                f"{CLAWDBOT_API_BASE}/command",
                json={
                    "command": f"send to {recipient} {message}",
                    "user_id": "mcp_agent",
                    "platform": target_platform,
                },
            )

            result = cmd_resp.json()

            return {
                "success": result.get("success", False),
                "message": result.get("message", ""),
                "recipient": contact_name,
                "platform": target_platform,
                "recipient_id": recipient_id,
                "data": result.get("data"),
            }

    except ImportError:
        return {"success": False, "error": "httpx not installed"}
    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_clawdbot_get_contacts(
    query: Optional[str] = None, limit: int = 10
) -> Dict[str, Any]:
    """List or search contacts."""
    try:
        import httpx

        async with httpx.AsyncClient(timeout=10.0) as client:
            if query:
                resp = await client.get(
                    f"{CLAWDBOT_API_BASE}/contacts/search",
                    params={"q": query, "limit": limit},
                )
            else:
                resp = await client.get(f"{CLAWDBOT_API_BASE}/contacts")

            if resp.status_code != 200:
                return {"success": False, "error": f"API error: {resp.status_code}"}

            data = resp.json()

            return {"success": True, "query": query, "contacts": data}

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_clawdbot_get_status() -> Dict[str, Any]:
    """Get Clawdbot bridge status."""
    try:
        import httpx

        async with httpx.AsyncClient(timeout=5.0) as client:
            status_resp = await client.get(f"{CLAWDBOT_API_BASE}/status")
            status_data = status_resp.json() if status_resp.status_code == 200 else {}

            sessions_resp = await client.get(f"{CLAWDBOT_API_BASE}/sessions")
            sessions_data = (
                sessions_resp.json() if sessions_resp.status_code == 200 else []
            )

            return {
                "success": True,
                "bridge_status": status_data.get("status", "unknown"),
                "initialized": status_data.get("initialized", False),
                "active_sessions": len(sessions_data),
                "sessions": sessions_data,
                "capabilities": status_data.get("capabilities", []),
            }

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_clawdbot_get_variables(
    include_templates: bool = True,
) -> Dict[str, Any]:
    """Get variables and templates."""
    try:
        import httpx

        async with httpx.AsyncClient(timeout=5.0) as client:
            var_resp = await client.get(f"{CLAWDBOT_API_BASE}/variables")
            variables = var_resp.json() if var_resp.status_code == 200 else {}

            result = {"success": True, "variables": variables}

            if include_templates:
                tmpl_resp = await client.get(f"{CLAWDBOT_API_BASE}/templates")
                result["templates"] = (
                    tmpl_resp.json() if tmpl_resp.status_code == 200 else {}
                )

            return result

    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_clawdbot_browser_open(url: str) -> Dict[str, Any]:
    """Open a URL in the browser via Clawdbot."""
    try:
        import httpx

        if not url.startswith("http://") and not url.startswith("https://"):
            url = f"https://{url}"

        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.post(
                f"{CLAWDBOT_API_BASE}/command",
                json={
                    "command": f"open {url}",
                    "user_id": "mcp_agent",
                    "platform": "browser",
                },
            )
            result = resp.json()

        return {
            "success": result.get("success", False),
            "message": result.get("message", ""),
            "url": url,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_clawdbot_browser_search(query: str) -> Dict[str, Any]:
    """Search the web via Clawdbot."""
    try:
        import httpx

        search_url = f"https://www.google.com/search?q={query.replace(' ', '+')}"

        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.post(
                f"{CLAWDBOT_API_BASE}/command",
                json={
                    "command": f"open {search_url}",
                    "user_id": "mcp_agent",
                    "platform": "browser",
                },
            )
            result = resp.json()

        return {
            "success": result.get("success", False),
            "message": result.get("message", ""),
            "query": query,
            "url": search_url,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_clawdbot_browser_read_page() -> Dict[str, Any]:
    """Read current browser page via screen OCR."""
    try:
        result = await handle_read_screen(monitor_id=0)
        result.pop("screenshot_base64", None)
        return {
            "success": result.get("success", False),
            "text": result.get("text", ""),
            "text_length": result.get("text_length", 0),
            "source": "screen_ocr",
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


async def handle_clawdbot_report_findings(
    findings: str,
    recipient: Optional[str] = None,
    platform: Optional[str] = None,
    title: Optional[str] = None,
) -> Dict[str, Any]:
    """Report findings via Clawdbot - either to a contact or via callback."""
    try:
        import httpx

        # Format the report message
        report_msg = ""
        if title:
            report_msg = f"📋 {title}\n\n"
        report_msg += findings

        # Truncate if too long
        if len(report_msg) > 4000:
            report_msg = report_msg[:3950] + "\n\n... [gekürzt]"

        if recipient:
            # Send to a specific contact
            result = await handle_clawdbot_send_message(
                recipient=recipient, message=report_msg, platform=platform
            )
            result["report_type"] = "contact_message"
            return result
        else:
            # Send via Clawdbot callback
            callback_payload = {
                "user_id": "mcp_agent",
                "platform": platform or "api",
                "success": True,
                "message": report_msg,
                "data": {
                    "type": "findings_report",
                    "title": title,
                    "findings_length": len(findings),
                },
            }

            async with httpx.AsyncClient(timeout=10.0) as client:
                # Try Clawdbot Gateway callback
                try:
                    resp = await client.post(
                        "http://localhost:18789/plugins/automation-ui/results",
                        json=callback_payload,
                    )
                    if resp.status_code == 200:
                        return {
                            "success": True,
                            "message": "Findings reported via Clawdbot callback",
                            "report_type": "callback",
                            "findings_length": len(findings),
                        }
                except Exception:
                    pass

                # Fallback: notify via Clawdbot API
                try:
                    resp = await client.post(
                        f"{CLAWDBOT_API_BASE}/notify",
                        params={
                            "user_id": "mcp_agent",
                            "platform": platform or "api",
                            "message": report_msg,
                            "notification_type": "info",
                        },
                    )
                    if resp.status_code == 200:
                        return {
                            "success": True,
                            "message": "Findings sent as notification",
                            "report_type": "notification",
                            "findings_length": len(findings),
                        }
                except Exception:
                    pass

                return {
                    "success": True,
                    "message": "Findings captured (no callback endpoint available)",
                    "report_type": "local",
                    "findings_preview": findings[:500],
                    "findings_length": len(findings),
                }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================
# Phase 1 Handlers — File/System, Documents, Clarify/Notify, Events, Smart Elements
# ============================================

# Phase 1.2: clarify / notify tools.
from agents.handoff.clarify_notify import handle_clarify  # noqa: E402
from agents.handoff.clarify_notify import handle_clarify_check, handle_notify
# Phase 1.5: document processing tools.
from agents.handoff.document_tools import handle_doc_apply  # noqa: E402
from agents.handoff.document_tools import (handle_doc_edit, handle_doc_export,
                                           handle_doc_list, handle_doc_scan)
# Phase 4.1: eyeTerm bridge tools.
from agents.handoff.eyeterm_bridge import (  # noqa: E402
    handle_collaborative_select, handle_eyeterm_calibrate,
    handle_eyeterm_dwell_request, handle_eyeterm_get_gaze,
    handle_eyeterm_status)
# Phase 1.3: file/system tools live in their own module.
from agents.handoff.file_system_tools import handle_dir_list  # noqa: E402
from agents.handoff.file_system_tools import (handle_file_open,
                                              handle_file_read,
                                              handle_file_search,
                                              handle_file_write,
                                              handle_process_kill,
                                              handle_process_list,
                                              handle_shell, handle_system_info)
# Phase 5: Geometry-first OCR (PaddleOCR > Tesseract+TSV > EasyOCR).
from agents.handoff.geometry_ocr import find_text_in_geometry  # noqa: E402
from agents.handoff.geometry_ocr import handle_read_screen_geometry
# Phase 6: Direct Office COM automation (replaces fragile keyboard simulation).
from agents.handoff.office_automation import handle_excel_fill  # noqa: E402
from agents.handoff.office_automation import (handle_excel_read,
                                              handle_office_save,
                                              handle_word_write)
# Phase 4.2 + 4.3: rich screen description + streaming subscription tools.
from agents.handoff.screen_description import (  # noqa: E402
    handle_describe_focus, handle_describe_screen, handle_get_cursor_context,
    handle_get_screen_changes, handle_get_selection, handle_get_window_tree,
    handle_list_actionable, handle_subscribe_screen_changes,
    handle_unsubscribe_screen_changes)

# Phase 7: VibeMind Bridge Tools moved to vibemind MCP (vibemind_mcp.py).
# See memory: project_vibemind_mcp_split.md


# ─── Phase 1.1: Event Queue tools ────────────────────────────────────────────
# These wrap the EventQueue datastructures (TaskEvent, active_tasks,
# completed_tasks) directly. We DO NOT start the EventQueue's background
# processing loops — the MCP only needs queue / lookup / cancel semantics,
# the actual execution path goes through handle_execute.


async def handle_event_add(
    event_type: str,
    payload: Optional[Dict[str, Any]] = None,
    priority: int = 0,
) -> Dict[str, Any]:
    """Queue a new event in the runtime EventQueue. Returns the event_id."""
    runtime = await get_runtime()
    eq = runtime.event_queue
    if eq is None:
        return {"success": False, "error": "EventQueue unavailable"}

    try:
        from core.event_queue import TaskEvent, TaskStatus  # type: ignore
    except Exception as e:
        return {"success": False, "error": f"event_queue types unavailable: {e}"}

    eq._task_counter += 1
    event = TaskEvent(
        id=f"task_{eq._task_counter}_{int(__import__('time').time())}",
        goal=event_type,
        context={"payload": payload or {}, "priority": int(priority)},
        status=TaskStatus.PENDING,
    )
    eq.active_tasks[event.id] = event

    return {
        "success": True,
        "event_id": event.id,
        "event_type": event_type,
        "status": event.status.value,
        "priority": int(priority),
    }


async def handle_event_status(event_id: str) -> Dict[str, Any]:
    """Look up an event by id."""
    runtime = await get_runtime()
    eq = runtime.event_queue
    if eq is None:
        return {"success": False, "error": "EventQueue unavailable"}

    event = eq.active_tasks.get(event_id)
    if event is None:
        for completed in eq.completed_tasks:
            if completed.id == event_id:
                event = completed
                break

    if event is None:
        return {"success": False, "error": f"unknown event_id: {event_id}"}

    return {
        "success": True,
        "event_id": event.id,
        "event_type": event.goal,
        "status": event.status.value,
        "created_at": event.created_at,
        "started_at": event.started_at,
        "completed_at": event.completed_at,
        "context": event.context,
        "actions_count": len(event.actions),
        "error": event.error,
    }


async def handle_event_list(
    status_filter: Optional[str] = None,
    limit: int = 50,
) -> Dict[str, Any]:
    """List events in the queue, optionally filtered by status."""
    runtime = await get_runtime()
    eq = runtime.event_queue
    if eq is None:
        return {"success": False, "error": "EventQueue unavailable"}

    all_events = list(eq.active_tasks.values()) + list(eq.completed_tasks)

    if status_filter:
        all_events = [e for e in all_events if e.status.value == status_filter]

    all_events = all_events[-limit:]

    return {
        "success": True,
        "count": len(all_events),
        "filter": status_filter,
        "events": [
            {
                "event_id": e.id,
                "event_type": e.goal,
                "status": e.status.value,
                "created_at": e.created_at,
                "actions_count": len(e.actions),
            }
            for e in all_events
        ],
    }


async def handle_event_cancel(event_id: str) -> Dict[str, Any]:
    """Cancel a queued/active event."""
    runtime = await get_runtime()
    eq = runtime.event_queue
    if eq is None:
        return {"success": False, "error": "EventQueue unavailable"}

    try:
        from core.event_queue import TaskStatus  # type: ignore
    except Exception as e:
        return {"success": False, "error": f"event_queue types unavailable: {e}"}

    event = eq.active_tasks.get(event_id)
    if event is None:
        return {"success": False, "error": f"unknown event_id: {event_id}"}

    if event.status in {TaskStatus.COMPLETED, TaskStatus.FAILED, TaskStatus.CANCELLED}:
        return {
            "success": False,
            "error": f"event already in terminal state: {event.status.value}",
        }

    event.status = TaskStatus.CANCELLED
    import time as _time

    event.completed_at = _time.time()
    return {"success": True, "event_id": event_id, "status": event.status.value}


async def handle_batch_execute(
    events: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """Bulk-add a set of events with a shared batch_id."""
    runtime = await get_runtime()
    eq = runtime.event_queue
    if eq is None:
        return {"success": False, "error": "EventQueue unavailable"}

    try:
        from core.event_queue import TaskEvent, TaskStatus  # type: ignore
    except Exception as e:
        return {"success": False, "error": f"event_queue types unavailable: {e}"}

    import time as _time
    import uuid as _uuid

    batch_id = f"batch_{_uuid.uuid4().hex[:10]}"
    created: List[str] = []

    for spec in events:
        event_type = spec.get("event_type") or spec.get("goal") or "unnamed"
        payload = spec.get("payload") or {}
        priority = int(spec.get("priority", 0))

        eq._task_counter += 1
        event = TaskEvent(
            id=f"task_{eq._task_counter}_{int(_time.time())}",
            goal=event_type,
            context={"payload": payload, "priority": priority, "batch_id": batch_id},
            status=TaskStatus.PENDING,
        )
        eq.active_tasks[event.id] = event
        created.append(event.id)

    return {
        "success": True,
        "batch_id": batch_id,
        "count": len(created),
        "event_ids": created,
    }


# ─── Phase 1.4: Smart Elements tools ─────────────────────────────────────────


async def handle_find_element(
    text: Optional[str] = None,
    element_type: Optional[str] = None,
    near_text: Optional[str] = None,
    monitor_id: int = 0,
) -> Dict[str, Any]:
    """Find a UI element by text/type using a 3-tier fallback strategy:

    1. Pattern Store cache (instant, from past successful finds)
    2. Validation Team + UIA (live control-tree walk)
    3. PaddleOCR geometry scan (pixel-level text detection with click coords)

    Tier 3 is the breakthrough: when the text is NOT in a UIA control
    (canvas, PDF, image, SVG, embedded chat, rendered email), PaddleOCR
    finds it as a pixel region and returns click-ready (x, y) coordinates.
    """
    target = text or near_text or ""
    if not target and not element_type:
        return {"success": False, "error": "text, near_text or element_type required"}

    runtime = await get_runtime()

    # Tier 1: Pattern Store cache (fast path).
    if runtime.pattern_store is not None and target:
        try:
            ps = runtime.pattern_store
            cached = None
            if hasattr(ps, "get_element"):
                cached = ps.get_element(target)
            elif hasattr(ps, "find"):
                cached = ps.find(target)
            if cached:
                return {
                    "success": True,
                    "source": "pattern_store",
                    "target": target,
                    "element": cached,
                }
        except Exception as e:
            logger.debug(f"pattern_store lookup failed: {e}")

    # Tier 2: Validation Team + UIA live control-tree walk.
    team = await get_validation_team()
    vt_result = await team.validate_element(target)
    vt_dict = vt_result if isinstance(vt_result, dict) else {"raw": vt_result}
    # Check if Tier 2 actually found coordinates.
    has_location = bool(
        vt_dict.get("element_location")
        and vt_dict["element_location"].get("x") is not None
    )
    if has_location:
        return {
            "success": True,
            "source": "validation_team",
            "target": target,
            "element_type": element_type,
            "near_text": near_text,
            **vt_dict,
        }

    # Tier 3: PaddleOCR geometry scan — pixel-level text detection.
    # This catches text in canvases, PDFs, images, SVGs, etc. that UIA
    # cannot see. Returns click-ready (x, y) from the polygon center.
    try:
        geom_result = await handle_read_screen_geometry(
            monitor_id=monitor_id,
            engine="auto",
            include_normalized=False,  # save response size
        )
        if geom_result.get("success"):
            match = find_text_in_geometry(geom_result, target, fuzzy=True)
            if match:
                return {
                    "success": True,
                    "source": "geometry_ocr",
                    "target": target,
                    "element_location": {"x": match["x"], "y": match["y"]},
                    "confidence": match.get("confidence", 0),
                    "bbox": match.get("bbox"),
                    "matched_text": match.get("text"),
                    "engine": geom_result.get("engine", {}).get("ocr_engine"),
                }
    except Exception as e:
        logger.debug(f"geometry_ocr fallback failed: {e}")

    # All three tiers failed.
    return {
        "success": True,
        "source": "validation_team",
        "target": target,
        "element_type": element_type,
        "near_text": near_text,
        "element_location": None,
        "note": "Element not found in UIA, Pattern Store, or pixel OCR",
        **vt_dict,
    }


async def handle_scroll_to(
    target: str,
    element_type: Optional[str] = None,
    then_click: bool = False,
    direction: str = "down",
    max_scrolls: int = 10,
) -> Dict[str, Any]:
    """Scroll until target element is visible, optionally click it."""
    if not target:
        return {"success": False, "error": "target is required"}

    import pyautogui

    for attempt in range(max_scrolls):
        # Try to find the element on screen.
        find_result = await handle_find_element(text=target, element_type=element_type)
        loc = (
            (find_result.get("element") or {}).get("element_location")
            if find_result.get("success")
            else None
        )
        if (
            loc is None
            and find_result.get("success")
            and find_result.get("element_location")
        ):
            loc = find_result.get("element_location")

        if loc and loc.get("x") is not None and loc.get("y") is not None:
            if then_click:
                pyautogui.click(int(loc["x"]), int(loc["y"]))
            return {
                "success": True,
                "target": target,
                "found_after_scrolls": attempt,
                "location": loc,
                "clicked": then_click,
            }

        # Not found yet — scroll one notch and try again.
        clicks = 3 if direction == "down" else -3
        pyautogui.scroll(-abs(clicks) if direction == "down" else abs(clicks))
        await asyncio.sleep(0.25)

    return {
        "success": False,
        "target": target,
        "error": f"element not found after {max_scrolls} scrolls",
    }


# Create MCP server
server = Server("handoff")


# ─── Phase 3.2: Audit Log ────────────────────────────────────────────────────

_AUDIT_LOG_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "logs", "mcp_audit.jsonl"
)
os.makedirs(os.path.dirname(_AUDIT_LOG_PATH), exist_ok=True)


def _audit_summarize_args(args: Dict[str, Any]) -> Dict[str, Any]:
    """Strip / truncate large fields from arguments before logging."""
    if not isinstance(args, dict):
        return {"_repr": repr(args)[:200]}
    summary: Dict[str, Any] = {}
    for k, v in args.items():
        if isinstance(v, str) and len(v) > 200:
            summary[k] = v[:200] + f"... ({len(v)} chars)"
        elif isinstance(v, list) and len(v) > 10:
            summary[k] = [v[0], "...", v[-1], f"({len(v)} items)"]
        elif isinstance(v, (str, int, float, bool)) or v is None:
            summary[k] = v
        else:
            summary[k] = type(v).__name__
    return summary


def _audit_summarize_result(result: Any) -> Dict[str, Any]:
    """Reduce a tool result to its bookkeeping fields."""
    if not isinstance(result, dict):
        return {"_repr": repr(result)[:200]}
    keep_keys = {
        "success",
        "error",
        "count",
        "steps_executed",
        "recovery_attempts",
        "replan_count",
        "final_strategy",
        "memory_task_id",
        "event_id",
        "document_id",
        "clarify_id",
        "delivered",
        "ok",
    }
    return {k: result[k] for k in keep_keys if k in result}


def _audit_write(entry: Dict[str, Any]) -> None:
    try:
        with open(_AUDIT_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, default=str) + "\n")
    except Exception as e:
        logger.debug(f"audit log write failed: {e}")


@server.list_tools()
async def list_tools() -> List[Tool]:
    """List available tools."""
    return TOOLS


@server.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[TextContent]:
    """Handle tool calls. Phase 3.2 wraps every dispatch in an audit-log entry."""
    _audit_started_at = datetime.now()
    _audit_t0 = _audit_started_at.timestamp()
    try:
        if name == "handoff_plan":
            result = await handle_plan(
                goal=arguments.get("goal", ""), context=arguments.get("context")
            )
        elif name == "handoff_execute":
            result = await handle_execute(
                plan=arguments.get("plan", []),
                speed_factor=float(arguments.get("speed_factor", 1.0)),
                goal=arguments.get("goal"),
                recovery=bool(arguments.get("recovery", False)),
            )
        elif name == "handoff_validate":
            result = await handle_validate(
                target=arguments.get("target", ""),
                expected_state=arguments.get("expected_state"),
            )
        elif name == "handoff_action":
            result = await handle_action(
                action_type=arguments.get("action_type", ""),
                params=arguments.get("params", {}),
            )
        elif name == "handoff_status":
            result = await handle_status()
        elif name == "handoff_read_screen":
            result = await handle_read_screen(
                region=arguments.get("region"),
                monitor_id=arguments.get("monitor_id", 0),
                include_screenshot=bool(arguments.get("include_screenshot", False)),
                with_vision=bool(arguments.get("with_vision", False)),
                vision_prompt=arguments.get("vision_prompt"),
                prefer_uia=bool(arguments.get("prefer_uia", True)),
                text_source=arguments.get("text_source", "auto"),
            )
        elif name == "handoff_get_focus":
            result = await handle_get_focus()
        elif name == "handoff_scroll":
            result = await handle_scroll(
                direction=arguments.get("direction", "down"),
                amount=arguments.get("amount", 3),
                x=arguments.get("x"),
                y=arguments.get("y"),
            )
        # Claude CLI Tools
        elif name == "claude_cli_run":
            result = await handle_claude_run(
                prompt=arguments.get("prompt", ""),
                skill=arguments.get("skill"),
                output_format=arguments.get("output_format", "text"),
            )
        elif name == "claude_cli_skill":
            result = await handle_claude_skill(
                skill_name=arguments.get("skill_name", ""),
                inputs=arguments.get("inputs", {}),
                ui_context=arguments.get("ui_context"),
            )
        elif name == "claude_cli_status":
            result = await handle_claude_status()
        # Vision Analysis Tool
        elif name == "vision_analyze":
            result = await handle_vision_analyze(
                prompt=arguments.get("prompt", ""),
                mode=arguments.get("mode", "custom"),
                json_output=arguments.get("json_output", True),
                monitor_id=arguments.get("monitor_id", 0),
            )
        # Clawdbot Messaging Tools
        elif name == "clawdbot_send_message":
            result = await handle_clawdbot_send_message(
                recipient=arguments.get("recipient", ""),
                message=arguments.get("message", ""),
                platform=arguments.get("platform"),
            )
        elif name == "clawdbot_get_contacts":
            result = await handle_clawdbot_get_contacts(
                query=arguments.get("query"), limit=arguments.get("limit", 10)
            )
        elif name == "clawdbot_get_status":
            result = await handle_clawdbot_get_status()
        elif name == "clawdbot_get_variables":
            result = await handle_clawdbot_get_variables(
                include_templates=arguments.get("include_templates", True)
            )
        # Clawdbot Browser Tools
        elif name == "clawdbot_browser_open":
            result = await handle_clawdbot_browser_open(url=arguments.get("url", ""))
        elif name == "clawdbot_browser_search":
            result = await handle_clawdbot_browser_search(
                query=arguments.get("query", "")
            )
        elif name == "clawdbot_browser_read_page":
            result = await handle_clawdbot_browser_read_page()
        elif name == "clawdbot_report_findings":
            result = await handle_clawdbot_report_findings(
                findings=arguments.get("findings", ""),
                recipient=arguments.get("recipient"),
                platform=arguments.get("platform"),
                title=arguments.get("title"),
            )

        # ─── Phase 1.1: Event Queue Tools ───────────────────────────────
        elif name == "handoff_event_add":
            result = await handle_event_add(
                event_type=arguments.get("event_type", ""),
                payload=arguments.get("payload"),
                priority=int(arguments.get("priority", 0)),
            )
        elif name == "handoff_event_status":
            result = await handle_event_status(event_id=arguments.get("event_id", ""))
        elif name == "handoff_event_list":
            result = await handle_event_list(
                status_filter=arguments.get("status_filter"),
                limit=int(arguments.get("limit", 50)),
            )
        elif name == "handoff_event_cancel":
            result = await handle_event_cancel(event_id=arguments.get("event_id", ""))
        elif name == "handoff_batch_execute":
            result = await handle_batch_execute(events=arguments.get("events", []))

        # ─── Phase 1.2: User Interaction Tools ──────────────────────────
        elif name == "handoff_notify":
            result = await handle_notify(
                message=arguments.get("message", ""),
                title=arguments.get("title"),
                level=arguments.get("level", "info"),
                metadata=arguments.get("metadata"),
            )
        elif name == "handoff_clarify":
            result = await handle_clarify(
                question=arguments.get("question", ""),
                options=arguments.get("options"),
                timeout_seconds=arguments.get("timeout_seconds", 300),
                metadata=arguments.get("metadata"),
            )
        elif name == "handoff_clarify_check":
            result = await handle_clarify_check(
                clarify_id=arguments.get("clarify_id", "")
            )

        # ─── Phase 1.3: File / System Tools ─────────────────────────────
        elif name == "handoff_shell":
            result = await handle_shell(
                command=arguments.get("command", ""),
                timeout=int(arguments.get("timeout", 30)),
                cwd=arguments.get("cwd"),
            )
        elif name == "handoff_file_search":
            result = await handle_file_search(
                pattern=arguments.get("pattern", ""),
                root=arguments.get("root"),
                max_results=int(arguments.get("max_results", 200)),
            )
        elif name == "handoff_file_open":
            result = await handle_file_open(path=arguments.get("path", ""))
        elif name == "handoff_dir_list":
            result = await handle_dir_list(path=arguments.get("path"))
        elif name == "handoff_file_read":
            result = await handle_file_read(
                path=arguments.get("path", ""),
                max_bytes=int(arguments.get("max_bytes", 1048576)),
                encoding=arguments.get("encoding", "utf-8"),
            )
        elif name == "handoff_file_write":
            result = await handle_file_write(
                path=arguments.get("path", ""),
                content=arguments.get("content", ""),
                mode=arguments.get("mode", "w"),
                encoding=arguments.get("encoding", "utf-8"),
                confirm=bool(arguments.get("confirm", False)),
                create_parents=bool(arguments.get("create_parents", False)),
            )
        elif name == "handoff_process_list":
            result = await handle_process_list(
                name_filter=arguments.get("name_filter"),
                limit=int(arguments.get("limit", 100)),
            )
        elif name == "handoff_process_kill":
            result = await handle_process_kill(
                pid=int(arguments.get("pid", -1)),
                force=bool(arguments.get("force", False)),
            )
        elif name == "handoff_system_info":
            result = await handle_system_info()

        # ─── Phase 1.4: Smart Elements Tools ────────────────────────────
        elif name == "handoff_find_element":
            result = await handle_find_element(
                text=arguments.get("text"),
                element_type=arguments.get("element_type"),
                near_text=arguments.get("near_text"),
            )
        elif name == "handoff_scroll_to":
            result = await handle_scroll_to(
                target=arguments.get("target", ""),
                element_type=arguments.get("element_type"),
                then_click=bool(arguments.get("then_click", False)),
                direction=arguments.get("direction", "down"),
                max_scrolls=int(arguments.get("max_scrolls", 10)),
            )

        # ─── Phase 1.5: Document Processing Tools ───────────────────────
        elif name == "handoff_doc_scan":
            result = await handle_doc_scan(
                path=arguments.get("path"),
                max_pages=int(arguments.get("max_pages", 20)),
                detect_structure=bool(arguments.get("detect_structure", True)),
            )
        elif name == "handoff_doc_edit":
            result = await handle_doc_edit(
                document_id=arguments.get("document_id", ""),
                page=int(arguments.get("page", 1)),
                section_index=int(arguments.get("section_index", 0)),
                new_text=arguments.get("new_text", ""),
                operation=arguments.get("operation", "replace"),
            )
        elif name == "handoff_doc_apply":
            result = await handle_doc_apply(
                document_id=arguments.get("document_id", ""),
                dry_run=bool(arguments.get("dry_run", False)),
            )
        elif name == "handoff_doc_export":
            result = await handle_doc_export(
                document_id=arguments.get("document_id", ""),
                format=arguments.get("format", "json"),
            )
        elif name == "handoff_doc_list":
            result = await handle_doc_list()

        # ─── Phase 4.1: eyeTerm Bridge ──────────────────────────────────
        elif name == "handoff_eyeterm_status":
            result = await handle_eyeterm_status()
        elif name == "handoff_eyeterm_get_gaze":
            result = await handle_eyeterm_get_gaze()
        elif name == "handoff_eyeterm_dwell_request":
            result = await handle_eyeterm_dwell_request(
                prompt=arguments.get("prompt", ""),
                dwell_ms=int(arguments.get("dwell_ms", 2000)),
                timeout_ms=int(arguments.get("timeout_ms", 15000)),
                confidence_threshold=float(arguments.get("confidence_threshold", 0.6)),
            )
        elif name == "handoff_eyeterm_calibrate":
            result = await handle_eyeterm_calibrate()
        elif name == "handoff_collaborative_select":
            result = await handle_collaborative_select(
                question=arguments.get("question", ""),
                mode=arguments.get("mode", "click"),
                timeout_seconds=int(arguments.get("timeout_seconds", 30)),
            )

        # ─── Phase 4.2: Screen Description ──────────────────────────────
        elif name == "handoff_describe_screen":
            result = await handle_describe_screen(
                detail=arguments.get("detail", "summary")
            )
        elif name == "handoff_describe_focus":
            result = await handle_describe_focus()
        elif name == "handoff_get_window_tree":
            result = await handle_get_window_tree()
        elif name == "handoff_get_selection":
            result = await handle_get_selection()
        elif name == "handoff_get_cursor_context":
            result = await handle_get_cursor_context()
        elif name == "handoff_list_actionable":
            result = await handle_list_actionable(
                window_hwnd=arguments.get("window_hwnd"),
                max_elements=int(arguments.get("max_elements", 100)),
            )

        # ─── Phase 4.3: Screen Change Subscription ──────────────────────
        elif name == "handoff_subscribe_screen_changes":
            result = await handle_subscribe_screen_changes(
                interval_ms=int(arguments.get("interval_ms", 500)),
                subscription_id=arguments.get("subscription_id"),
            )
        elif name == "handoff_unsubscribe_screen_changes":
            result = await handle_unsubscribe_screen_changes(
                subscription_id=arguments.get("subscription_id", ""),
            )
        elif name == "handoff_get_screen_changes":
            result = await handle_get_screen_changes(
                subscription_id=arguments.get("subscription_id", ""),
            )

        # ─── Phase 8: Self-Awareness ────────────────────────────────────
        elif name == "handoff_self_info":
            result = {
                "success": True,
                "server": "handoff",
                "tool_count": len(TOOLS),
                "python_version": sys.version.split()[0],
                "pid": os.getpid(),
                "uptime_note": "since last MCP restart",
                "hub": (await get_runtime()).get_hub_status(),
            }

        # Phase 7 VibeMind tools moved to vibemind MCP

        # ─── Phase 6: Office COM Automation ─────────────────────────────
        elif name == "handoff_excel_fill":
            result = await handle_excel_fill(
                cells=arguments.get("cells", []),
                auto_fit=bool(arguments.get("auto_fit", True)),
                sheet_name=arguments.get("sheet_name"),
            )
        elif name == "handoff_excel_read":
            result = await handle_excel_read(
                range_ref=arguments.get("range_ref", "A1"),
                sheet_name=arguments.get("sheet_name"),
            )
        elif name == "handoff_word_write":
            result = await handle_word_write(
                paragraphs=arguments.get("paragraphs", []),
                clear_first=bool(arguments.get("clear_first", False)),
            )
        elif name == "handoff_office_save":
            result = await handle_office_save(
                app=arguments.get("app", "auto"),
            )

        # ─── Phase 5: Geometry OCR ──────────────────────────────────────
        elif name == "handoff_read_screen_geometry":
            result = await handle_read_screen_geometry(
                monitor_id=int(arguments.get("monitor_id", 0)),
                region=arguments.get("region"),
                engine=arguments.get("engine", "auto"),
                return_char_boxes=bool(arguments.get("return_char_boxes", False)),
                include_normalized=bool(arguments.get("include_normalized", True)),
            )

        else:
            result = {"error": f"Unknown tool: {name}"}

        # Phase 3.2: emit audit log entry on success.
        _audit_write(
            {
                "ts": _audit_started_at.isoformat(),
                "tool": name,
                "args": _audit_summarize_args(arguments),
                "result": _audit_summarize_result(result),
                "duration_ms": round(
                    (datetime.now().timestamp() - _audit_t0) * 1000, 2
                ),
                "ok": True,
            }
        )

        return [TextContent(type="text", text=json.dumps(result, indent=2))]

    except Exception as e:
        # Phase 3.2: emit audit log entry on failure.
        _audit_write(
            {
                "ts": _audit_started_at.isoformat(),
                "tool": name,
                "args": _audit_summarize_args(arguments),
                "error": str(e),
                "duration_ms": round(
                    (datetime.now().timestamp() - _audit_t0) * 1000, 2
                ),
                "ok": False,
            }
        )
        return [TextContent(type="text", text=json.dumps({"error": str(e)}))]


async def cleanup():
    """Clean up resources on shutdown."""
    global _planning_team, _validation_team, _runtime

    logger.info("Starting cleanup...")

    if _planning_team:
        if hasattr(_planning_team, "llm_client") and _planning_team.llm_client:
            await _planning_team.llm_client.close()
        await _planning_team.stop()
        _planning_team = None
        logger.info("Planning team stopped")

    if _validation_team:
        await _validation_team.stop()
        _validation_team = None
        logger.info("Validation team stopped")

    if _runtime:
        await _runtime.stop()
        _runtime = None
        logger.info("Runtime stopped")

    logger.info("Cleanup complete")


# Global shutdown flag
_shutdown_requested = False


def signal_handler(signum, frame):
    """Handle shutdown signals for graceful termination."""
    global _shutdown_requested
    signal_name = signal.Signals(signum).name
    logger.info(f"Received signal {signal_name}, initiating graceful shutdown...")
    _shutdown_requested = True


def setup_signal_handlers():
    """Setup signal handlers for graceful shutdown."""
    if sys.platform == "win32":
        # Windows: Handle SIGINT (Ctrl+C) and SIGBREAK (Ctrl+Break)
        signal.signal(signal.SIGINT, signal_handler)
        try:
            signal.signal(signal.SIGBREAK, signal_handler)
        except AttributeError:
            pass  # SIGBREAK not available on all platforms
    else:
        # Unix: Handle SIGTERM and SIGINT
        signal.signal(signal.SIGTERM, signal_handler)
        signal.signal(signal.SIGINT, signal_handler)


async def main():
    """Run the MCP server."""
    # Setup signal handlers
    setup_signal_handlers()

    # Log startup
    logger.info("=" * 50)
    logger.info("Handoff MCP Server starting...")
    logger.info(f"Python: {sys.version}")
    logger.info(f"Platform: {sys.platform}")
    if _config:
        logger.info(
            f"Config loaded: MoireServer {_config.moire_host}:{_config.moire_port}"
        )
    logger.info("=" * 50)

    # Record start time
    start_time = datetime.now()

    try:
        async with stdio_server() as (read_stream, write_stream):
            logger.info("MCP stdio server started, ready for connections")
            await server.run(
                read_stream, write_stream, server.create_initialization_options()
            )
    except KeyboardInterrupt:
        logger.info("Keyboard interrupt received")
    except Exception as e:
        logger.error(f"Server error: {e}")
        raise
    finally:
        # Log uptime
        uptime = datetime.now() - start_time
        logger.info(f"Server uptime: {uptime}")

        # Cleanup
        await cleanup()
        logger.info("Handoff MCP Server shutdown complete")


# ─── Phase: app lifecycle + adaptive skill library (4-way-split perception) ───
# These 7 SYNC handlers are called by name from the perception/document
# entrypoints (which import this module as `H`). They MUST stay synchronous —
# they are dispatched from a `_SYNC` table — so any async helper from
# window_focus is driven through a fresh event loop via `_run_sync` below.

# Friendly app-name → executable map. Anything not listed is treated as a raw
# path / command and handed to subprocess.Popen / os.startfile directly.
_APP_ALIASES = {
    "notepad": "notepad.exe",
    "calc": "calc.exe",
    "calculator": "calc.exe",
    "explorer": "explorer.exe",
    "cmd": "cmd.exe",
    "powershell": "powershell.exe",
    "wordpad": "wordpad.exe",
    "paint": "mspaint.exe",
    "mspaint": "mspaint.exe",
}


def _run_sync(coro):
    """Drive an async coroutine to completion from sync code.

    The window_focus helpers (focus_window, get_active_window) are async, but
    these handlers are sync. Use a private event loop so we never collide with
    a loop that may already be running on the calling thread.
    """
    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()


def _app_launch(app_name: str, args=None) -> dict:
    """Launch an app/executable by friendly name or full path.

    Resolves common aliases (notepad/calc/explorer/cmd/...) and otherwise treats
    ``app_name`` as a path or command. ``args`` may be a list or a single string.
    """
    import subprocess

    try:
        if not app_name or not str(app_name).strip():
            return {"success": False, "app": app_name, "pid": None, "error": "app_name is empty"}

        target = _APP_ALIASES.get(str(app_name).strip().lower(), str(app_name).strip())

        # Normalize args into a list of strings.
        if args is None:
            arg_list: List[str] = []
        elif isinstance(args, str):
            arg_list = [args] if args else []
        elif isinstance(args, (list, tuple)):
            arg_list = [str(a) for a in args]
        else:
            arg_list = [str(args)]

        try:
            proc = subprocess.Popen([target] + arg_list)
            return {"success": True, "app": target, "pid": proc.pid, "error": None}
        except (FileNotFoundError, OSError) as popen_err:
            # Fall back to os.startfile for documents / shell-resolved names when
            # there are no extra args (startfile cannot pass argv).
            if not arg_list:
                try:
                    os.startfile(target)  # type: ignore[attr-defined]  # Windows-only
                    return {"success": True, "app": target, "pid": None, "error": None}
                except Exception as start_err:
                    return {
                        "success": False,
                        "app": target,
                        "pid": None,
                        "error": f"{popen_err} / startfile: {start_err}",
                    }
            return {"success": False, "app": target, "pid": None, "error": str(popen_err)}
    except Exception as e:
        logger.error(f"_app_launch failed: {e}")
        return {"success": False, "app": app_name, "pid": None, "error": str(e)}


def _app_focus(title_substring: str) -> dict:
    """Focus an existing visible window whose title contains ``title_substring``."""
    try:
        from agents.handoff.window_focus import find_window_by_title, focus_window

        if not title_substring or not str(title_substring).strip():
            return {"success": False, "error": "title_substring is empty", "title": None, "hwnd": None}

        hwnd = find_window_by_title(str(title_substring), exact=False)
        if not hwnd:
            return {
                "success": False,
                "error": f"No visible window matching '{title_substring}'",
                "title": None,
                "hwnd": None,
            }

        focused = _run_sync(focus_window(hwnd))
        return {
            "success": bool(focused),
            "hwnd": hwnd,
            "title": title_substring,
            "error": None if focused else "SetForegroundWindow returned false",
        }
    except Exception as e:
        logger.error(f"_app_focus failed: {e}")
        return {"success": False, "error": str(e), "title": None, "hwnd": None}


def _app_list_running() -> dict:
    """List visible top-level windows (title + hwnd, plus pid when resolvable)."""
    try:
        import ctypes

        from agents.handoff.window_focus import list_visible_windows

        windows = list_visible_windows()
        user32 = ctypes.windll.user32
        for w in windows:
            try:
                pid = ctypes.c_ulong()
                user32.GetWindowThreadProcessId(int(w["hwnd"]), ctypes.byref(pid))
                w["pid"] = pid.value or None
            except Exception:
                w["pid"] = None
        return {"success": True, "windows": windows}
    except Exception as e:
        logger.error(f"_app_list_running failed: {e}")
        return {"success": False, "windows": [], "error": str(e)}


def _app_launch_or_focus(app_name: str, title_hint=None, args=None) -> dict:
    """Focus a matching window if one exists, otherwise launch the app.

    Matching is tried first against ``title_hint`` then against ``app_name``.
    """
    try:
        from agents.handoff.window_focus import find_window_by_title

        for candidate in (title_hint, app_name):
            if candidate and str(candidate).strip():
                hwnd = find_window_by_title(str(candidate), exact=False)
                if hwnd:
                    result = _app_focus(str(candidate))
                    result["action"] = "focused"
                    return result

        result = _app_launch(app_name, args=args)
        result["action"] = "launched"
        return result
    except Exception as e:
        logger.error(f"_app_launch_or_focus failed: {e}")
        return {"success": False, "app": app_name, "pid": None, "error": str(e)}


def _iter_skill_files():
    """Yield (app, skill_name, path) for every SKILL.md under _SKILL_LIB_ROOT."""
    root = _SKILL_LIB_ROOT
    if not os.path.isdir(root):
        return
    for dirpath, _dirnames, filenames in os.walk(root):
        for fn in filenames:
            if fn.lower() == "skill.md":
                path = os.path.join(dirpath, fn)
                rel = os.path.relpath(dirpath, root)
                parts = rel.split(os.sep) if rel != "." else []
                app = parts[0] if parts else ""
                skill_name = parts[-1] if parts else os.path.basename(dirpath)
                yield app, skill_name, path


def _read_skill_meta(path: str) -> dict:
    """Pull a one-line description (frontmatter or first prose line) + raw body."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception:
        return {"description": "", "body": ""}

    description = ""
    body = text
    # Minimal YAML-frontmatter parse (no external dep): grab `description:` / `name:`.
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            fm = text[3:end]
            body = text[end + 4 :]
            for line in fm.splitlines():
                ls = line.strip()
                if ls.lower().startswith("description:"):
                    description = ls.split(":", 1)[1].strip().strip("\"'")
                    break
    if not description:
        for line in body.splitlines():
            ls = line.strip().lstrip("#").strip()
            if ls:
                description = ls[:200]
                break
    return {"description": description, "body": body}


def _skill_search(query: str, agent=None, limit: int = 5) -> dict:
    """Search the adaptive-skill library for SKILL.md files matching ``query``.

    Dependency-light ranking: case-insensitive substring hits (weighted by field)
    plus token-overlap. No vector DB. ``agent`` optionally scopes to a subdir.
    """
    try:
        q = (query or "").strip().lower()
        if not q:
            return {"success": True, "results": []}
        q_tokens = set(t for t in q.replace("/", " ").replace("-", " ").split() if t)

        results = []
        for app, skill_name, path in _iter_skill_files():
            if agent and str(agent).strip().lower() not in (app.lower(), skill_name.lower()):
                continue
            meta = _read_skill_meta(path)
            name_l = skill_name.lower()
            desc_l = meta["description"].lower()
            body_l = meta["body"].lower()

            score = 0.0
            if q in name_l:
                score += 10.0
            if q in desc_l:
                score += 5.0
            if q in body_l:
                score += 2.0
            # Token overlap across name + description + body.
            field_tokens = set(
                t for t in (name_l + " " + desc_l + " " + body_l)
                .replace("/", " ").replace("-", " ").split()
                if t
            )
            score += 1.0 * len(q_tokens & field_tokens)

            if score > 0:
                results.append(
                    {
                        "app": app,
                        "skill_name": skill_name,
                        "path": path,
                        "score": round(score, 3),
                        "description": meta["description"],
                    }
                )

        results.sort(key=lambda r: r["score"], reverse=True)
        try:
            lim = int(limit)
        except (TypeError, ValueError):
            lim = 5
        return {"success": True, "results": results[: max(0, lim)]}
    except Exception as e:
        logger.error(f"_skill_search failed: {e}")
        return {"success": False, "results": [], "error": str(e)}


def _skill_list(app=None) -> dict:
    """List all SKILL.md skills, optionally filtered to a subdir/app."""
    try:
        skills = []
        flt = str(app).strip().lower() if app and str(app).strip() else None
        for skill_app, skill_name, path in _iter_skill_files():
            if flt and flt not in (skill_app.lower(), skill_name.lower()):
                continue
            meta = _read_skill_meta(path)
            skills.append(
                {
                    "app": skill_app,
                    "skill_name": skill_name,
                    "path": path,
                    "description": meta["description"],
                }
            )
        skills.sort(key=lambda s: (s["app"], s["skill_name"]))
        return {"success": True, "skills": skills}
    except Exception as e:
        logger.error(f"_skill_list failed: {e}")
        return {"success": False, "skills": [], "error": str(e)}


def _skill_save_and_index(app: str, skill_name: str, frontmatter: dict, body: str) -> dict:
    """Write skills/<app>/<skill_name>/SKILL.md with YAML frontmatter + body."""
    try:
        if not app or not str(app).strip():
            return {"success": False, "error": "app is empty", "path": None}
        if not skill_name or not str(skill_name).strip():
            return {"success": False, "error": "skill_name is empty", "path": None}

        skill_dir = os.path.join(_SKILL_LIB_ROOT, str(app).strip(), str(skill_name).strip())
        os.makedirs(skill_dir, exist_ok=True)
        path = os.path.join(skill_dir, "SKILL.md")

        def _yaml_scalar(v) -> str:
            if isinstance(v, bool):
                return "true" if v else "false"
            if isinstance(v, (int, float)):
                return str(v)
            if isinstance(v, (list, tuple)):
                return "[" + ", ".join(_yaml_scalar(x) for x in v) + "]"
            s = str(v)
            if s == "" or any(c in s for c in ':#\n"\'') or s.strip() != s:
                return '"' + s.replace("\\", "\\\\").replace('"', '\\"') + '"'
            return s

        lines = ["---"]
        for k, v in (frontmatter or {}).items():
            lines.append(f"{k}: {_yaml_scalar(v)}")
        lines.append("---")
        lines.append("")
        content = "\n".join(lines) + (body or "")

        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        return {"success": True, "path": path}
    except Exception as e:
        logger.error(f"_skill_save_and_index failed: {e}")
        return {"success": False, "error": str(e), "path": None}


# ─── Tool() schemas for the 7 app/skill handlers above ───
TOOLS += [
    Tool(
        name="app_launch",
        description="Launch an app/executable by friendly name (notepad, calc, explorer, cmd, ...) or full path/command.",
        inputSchema={
            "type": "object",
            "properties": {
                "app_name": {
                    "type": "string",
                    "description": "App alias (notepad, calc, explorer, cmd) or a full executable path/command",
                },
                "args": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional command-line arguments to pass to the executable",
                },
            },
            "required": ["app_name"],
        },
    ),
    Tool(
        name="app_focus",
        description="Focus an existing visible window whose title contains the given substring.",
        inputSchema={
            "type": "object",
            "properties": {
                "title_substring": {
                    "type": "string",
                    "description": "Substring to match against window titles (case-insensitive)",
                },
            },
            "required": ["title_substring"],
        },
    ),
    Tool(
        name="app_list_running",
        description="List visible top-level windows (title, hwnd, and pid when resolvable).",
        inputSchema={
            "type": "object",
            "properties": {},
        },
    ),
    Tool(
        name="app_launch_or_focus",
        description="Focus a matching window if one already exists (by title_hint or app_name), otherwise launch the app.",
        inputSchema={
            "type": "object",
            "properties": {
                "app_name": {
                    "type": "string",
                    "description": "App alias or full path/command to launch if no window matches",
                },
                "title_hint": {
                    "type": "string",
                    "description": "Optional window-title substring to look for before launching",
                },
                "args": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional command-line arguments used only when launching",
                },
            },
            "required": ["app_name"],
        },
    ),
    Tool(
        name="skill_search",
        description="Search the adaptive-skill library (SKILL.md tree) for skills matching a query. Ranks by name/description/body relevance; returns top results.",
        inputSchema={
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query matched against skill name, description, and body",
                },
                "agent": {
                    "type": "string",
                    "description": "Optional app/agent subdir to scope the search to",
                },
                "limit": {
                    "type": "integer",
                    "default": 5,
                    "description": "Maximum number of results to return",
                },
            },
            "required": ["query"],
        },
    ),
    Tool(
        name="skill_list",
        description="List available skills (all SKILL.md under the skill library), optionally filtered to a subdir/app.",
        inputSchema={
            "type": "object",
            "properties": {
                "app": {
                    "type": "string",
                    "description": "Optional app/subdir name to filter the listing",
                },
            },
        },
    ),
    Tool(
        name="skill_save_and_index",
        description="Write a new skills/<app>/<skill_name>/SKILL.md with YAML frontmatter (from a dict) plus a markdown body, creating directories as needed.",
        inputSchema={
            "type": "object",
            "properties": {
                "app": {
                    "type": "string",
                    "description": "App/subdir the skill belongs to",
                },
                "skill_name": {
                    "type": "string",
                    "description": "Name of the skill (used as the leaf directory)",
                },
                "frontmatter": {
                    "type": "object",
                    "description": "Key/value pairs emitted as YAML frontmatter",
                },
                "body": {
                    "type": "string",
                    "description": "Markdown body written after the frontmatter",
                },
            },
            "required": ["app", "skill_name", "frontmatter", "body"],
        },
    ),
]


# ─── Phase: data / knowledge-interpretation (4-way-split data server) ───
#
# Two distinct knowledge backends:
#   * ROWBOAT  — the local Rowboat server (HTTP, semantic KB + chat agent).
#   * ROARBOOT — a plain markdown vault on disk (no server; we read .md files).
#
# Plus a generic file-quality LLM evaluator. All functions here are SYNC
# (the data entrypoint dispatches them synchronously) and use `requests`.
# Every function is defensive: it returns a clean error dict instead of
# raising, because the data server treats a raised exception as a hard fault.

import glob as _glob

_ROWBOAT_URL = os.environ.get("ROWBOAT_URL", "http://localhost:3100").rstrip("/")
_ROWBOAT_PROJECT_ID = os.environ.get(
    "ROWBOAT_PROJECT_ID", "c157ade4-ebce-4d1b-a7a5-5fd70d238f8d"
)
_ROWBOAT_API_KEY = os.environ.get("ROWBOAT_API_KEY", "vibemind-local-key")
_ROARBOOT_ROOT = os.environ.get("ROARBOOT_ROOT", r"C:\Users\User\.rowboat\knowledge")


def _rowboat_headers() -> Dict[str, str]:
    return {"Authorization": f"Bearer {_ROWBOAT_API_KEY}"}


# ─── ROWBOAT (HTTP client to the local Rowboat server) ───


def _rowboat_chat(
    message: str,
    context: str = "default",
    conversation_id=None,
    timeout: int = 120,
) -> dict:
    """Chat with the Rowboat agent. VERIFIED contract:

    POST {ROWBOAT_URL}/api/v1/{projectId}/chat
      header Authorization: Bearer {key}
      body   {"messages":[{"role":"user","content": ...}]}
      resp   {"conversationId":..., "turn":{"output":[{role,content}...]}}

    The answer is the last assistant message in turn.output.
    """
    try:
        import requests
    except Exception as e:  # pragma: no cover - requests is in the venv
        return {"success": False, "error": f"requests unavailable: {e}"}

    url = f"{_ROWBOAT_URL}/api/v1/{_ROWBOAT_PROJECT_ID}/chat"
    body: Dict[str, Any] = {"messages": [{"role": "user", "content": message}]}
    if conversation_id:
        body["conversationId"] = conversation_id
    if context and context != "default":
        body["context"] = context
    try:
        resp = requests.post(
            url, json=body, headers=_rowboat_headers(), timeout=timeout
        )
    except Exception as e:
        return {"success": False, "error": f"rowboat request failed: {e}"}

    if not (200 <= resp.status_code < 300):
        return {
            "success": False,
            "error": f"rowboat chat HTTP {resp.status_code}",
            "status": resp.status_code,
            "body": resp.text[:500],
        }
    try:
        data = resp.json()
    except Exception as e:
        return {"success": False, "error": f"rowboat returned non-JSON: {e}"}

    answer = ""
    try:
        output = (data.get("turn") or {}).get("output") or []
        for msg in reversed(output):
            if isinstance(msg, dict) and msg.get("role") == "assistant":
                content = msg.get("content")
                answer = content if isinstance(content, str) else str(content)
                break
        if not answer and output:
            last = output[-1]
            if isinstance(last, dict):
                answer = str(last.get("content", ""))
    except Exception as e:
        return {"success": False, "error": f"rowboat parse error: {e}", "raw": data}

    return {
        "success": True,
        "answer": answer,
        "conversation_id": data.get("conversationId"),
        "raw": data,
    }


def _rowboat_upload(file_path: str, title=None, tags=None) -> dict:
    """Upload a knowledge document to Rowboat.

    The upload endpoint of the current Rowboat version is NOT verified, so this
    is best-effort: POST multipart to {ROWBOAT_URL}/api/v1/{projectId}/sources.
    On any non-2xx (or error) we return a clean dict with a hint — never raise.
    """
    try:
        import requests
    except Exception as e:  # pragma: no cover
        return {"success": False, "error": f"requests unavailable: {e}"}

    if not os.path.isfile(file_path):
        return {"success": False, "error": f"file not found: {file_path}"}

    url = f"{_ROWBOAT_URL}/api/v1/{_ROWBOAT_PROJECT_ID}/sources"
    data: Dict[str, Any] = {}
    if title:
        data["title"] = title
    if tags:
        data["tags"] = (
            ",".join(tags) if isinstance(tags, (list, tuple)) else str(tags)
        )
    try:
        with open(file_path, "rb") as fh:
            files = {"file": (os.path.basename(file_path), fh)}
            resp = requests.post(
                url,
                files=files,
                data=data,
                headers=_rowboat_headers(),
                timeout=120,
            )
    except Exception as e:
        return {
            "success": False,
            "error": f"rowboat upload failed: {e}",
            "hint": "upload endpoint may differ in this Rowboat version",
        }

    if not (200 <= resp.status_code < 300):
        return {
            "success": False,
            "error": f"rowboat upload HTTP {resp.status_code}",
            "status": resp.status_code,
            "hint": "upload endpoint may differ in this Rowboat version",
            "body": resp.text[:500],
        }
    try:
        payload = resp.json()
    except Exception:
        payload = {"text": resp.text[:500]}
    return {"success": True, "result": payload, "status": resp.status_code}


def _rowboat_search(query: str, folder=None, limit: int = 20) -> dict:
    """Semantic search over the Rowboat knowledge base.

    The search endpoint is NOT verified for this Rowboat version; best-effort
    POST {ROWBOAT_URL}/api/v1/{projectId}/search. Clean error on failure.
    """
    try:
        import requests
    except Exception as e:  # pragma: no cover
        return {"success": False, "error": f"requests unavailable: {e}"}

    url = f"{_ROWBOAT_URL}/api/v1/{_ROWBOAT_PROJECT_ID}/search"
    body: Dict[str, Any] = {"query": query, "limit": limit}
    if folder:
        body["folder"] = folder
    try:
        resp = requests.post(
            url, json=body, headers=_rowboat_headers(), timeout=60
        )
    except Exception as e:
        return {
            "success": False,
            "error": f"rowboat search failed: {e}",
            "hint": "search endpoint may differ in this Rowboat version",
        }

    if not (200 <= resp.status_code < 300):
        return {
            "success": False,
            "error": f"rowboat search HTTP {resp.status_code}",
            "status": resp.status_code,
            "hint": "search endpoint may differ in this Rowboat version",
            "body": resp.text[:500],
        }
    try:
        return {"success": True, "results": resp.json()}
    except Exception:
        return {"success": True, "results": resp.text[:2000]}


def _rowboat_list_folders(limit: int = 50) -> dict:
    """List Rowboat knowledge folders/sources.

    Best-effort GET {ROWBOAT_URL}/api/v1/{projectId}/sources. Clean error on
    failure (endpoint not verified for this Rowboat version).
    """
    try:
        import requests
    except Exception as e:  # pragma: no cover
        return {"success": False, "error": f"requests unavailable: {e}"}

    url = f"{_ROWBOAT_URL}/api/v1/{_ROWBOAT_PROJECT_ID}/sources"
    try:
        resp = requests.get(
            url,
            params={"limit": limit},
            headers=_rowboat_headers(),
            timeout=30,
        )
    except Exception as e:
        return {
            "success": False,
            "error": f"rowboat list_folders failed: {e}",
            "hint": "sources endpoint may differ in this Rowboat version",
        }

    if not (200 <= resp.status_code < 300):
        return {
            "success": False,
            "error": f"rowboat list_folders HTTP {resp.status_code}",
            "status": resp.status_code,
            "hint": "sources endpoint may differ in this Rowboat version",
            "body": resp.text[:500],
        }
    try:
        return {"success": True, "sources": resp.json()}
    except Exception:
        return {"success": True, "sources": resp.text[:2000]}


# ─── ROARBOOT (local markdown vault — filesystem, NOT Rowboat) ───


def _roarboot_resolve_root(root=None) -> str:
    return os.path.abspath(root or _ROARBOOT_ROOT)


def _roarboot_iter_md(base: str):
    """Yield .md file paths recursively under base."""
    if not os.path.isdir(base):
        return
    for path in _glob.glob(os.path.join(base, "**", "*.md"), recursive=True):
        if os.path.isfile(path):
            yield path


def _roarboot_list_folders(root=None) -> dict:
    """List immediate subfolders under the knowledge root."""
    base = _roarboot_resolve_root(root)
    if not os.path.isdir(base):
        return {"success": False, "error": f"root not found: {base}", "root": base}
    try:
        folders = sorted(
            name
            for name in os.listdir(base)
            if os.path.isdir(os.path.join(base, name))
        )
    except Exception as e:
        return {"success": False, "error": str(e), "root": base}
    return {"success": True, "folders": folders, "root": base}


def _roarboot_read_knowledge(
    folder: str,
    query=None,
    limit: int = 20,
    name_pattern=None,
    root=None,
) -> dict:
    """Read .md files in <root>/<folder> recursively.

    Optionally filter by name_pattern (glob or substring against the filename)
    and/or by query (case-insensitive substring against file content).
    Caps the result at `limit` files.
    """
    base = _roarboot_resolve_root(root)
    target = os.path.join(base, folder) if folder else base
    if not os.path.isdir(target):
        return {
            "success": False,
            "error": f"folder not found: {target}",
            "files": [],
            "count": 0,
        }

    import fnmatch

    q = (query or "").lower().strip()
    pat = name_pattern or ""
    has_glob = any(c in pat for c in "*?[") if pat else False

    files: List[Dict[str, str]] = []
    for path in _roarboot_iter_md(target):
        name = os.path.basename(path)
        if pat:
            if has_glob:
                if not fnmatch.fnmatch(name.lower(), pat.lower()):
                    continue
            elif pat.lower() not in name.lower():
                continue
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as fh:
                text = fh.read()
        except Exception:
            continue
        if q and q not in text.lower():
            continue
        excerpt = text[:500].strip()
        files.append({"path": path, "name": name, "excerpt": excerpt})
        if len(files) >= max(1, int(limit)):
            break

    return {"success": True, "files": files, "count": len(files)}


# ─── Minimal LLM client (OpenAI chat completions via requests) ───
#
# This file has no general-purpose text-chat LLM client — `vision_agent` is
# image-only and `_get_st_model` is an embedder. So we implement a small,
# dependency-light OpenAI chat client here (requests POST). If no
# OPENAI_API_KEY is set we degrade gracefully (callers report which files
# they would have used).


def _llm_chat_complete(
    messages: List[Dict[str, str]], model: str = "gpt-4o-mini", timeout: int = 90
) -> dict:
    """Call OpenAI chat/completions. Returns {"success":bool,"content":str|None,...}."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return {"success": False, "error": "no OPENAI_API_KEY"}
    try:
        import requests
    except Exception as e:  # pragma: no cover
        return {"success": False, "error": f"requests unavailable: {e}"}
    try:
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={"model": model, "messages": messages, "temperature": 0.2},
            timeout=timeout,
        )
    except Exception as e:
        return {"success": False, "error": f"llm request failed: {e}"}
    if not (200 <= resp.status_code < 300):
        return {
            "success": False,
            "error": f"llm HTTP {resp.status_code}",
            "body": resp.text[:500],
        }
    try:
        data = resp.json()
        content = data["choices"][0]["message"]["content"]
    except Exception as e:
        return {"success": False, "error": f"llm parse error: {e}"}
    return {"success": True, "content": content}


def _tokenize(text: str) -> set:
    import re

    return {t for t in re.findall(r"[a-z0-9]{3,}", (text or "").lower())}


def _roarboot_ask(
    question: str,
    folder=None,
    max_files: int = 10,
    max_chars_per_file: int = 4000,
    model: str = "gpt-4o-mini",
    root=None,
) -> dict:
    """RAG over the markdown vault: rank .md files by overlap with the question,
    build a context prompt, and ask the LLM. Graceful no-key fallback."""
    base = _roarboot_resolve_root(root)
    target = os.path.join(base, folder) if folder else base
    if not os.path.isdir(target):
        return {"success": False, "error": f"folder not found: {target}"}

    q_tokens = _tokenize(question)
    scored: List[tuple] = []
    for path in _roarboot_iter_md(target):
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as fh:
                text = fh.read()
        except Exception:
            continue
        overlap = len(q_tokens & _tokenize(text)) if q_tokens else 0
        scored.append((overlap, path, text))

    if not scored:
        return {"success": False, "error": f"no .md files under {target}"}

    # Rank by token overlap (desc); ties keep filesystem order.
    scored.sort(key=lambda r: r[0], reverse=True)
    chosen = scored[: max(1, int(max_files))]
    sources = [p for _, p, _ in chosen]

    context_blocks = []
    for _, path, text in chosen:
        snippet = text[: max(200, int(max_chars_per_file))]
        context_blocks.append(f"### FILE: {path}\n{snippet}")
    context = "\n\n".join(context_blocks)

    prompt = (
        "Answer the question using ONLY the provided knowledge files. "
        "Cite file names where relevant. If the answer is not present, say so.\n\n"
        f"QUESTION:\n{question}\n\nKNOWLEDGE:\n{context}"
    )
    llm = _llm_chat_complete(
        [
            {"role": "system", "content": "You are a precise knowledge assistant."},
            {"role": "user", "content": prompt},
        ],
        model=model,
    )
    if not llm.get("success"):
        if llm.get("error") == "no OPENAI_API_KEY":
            return {
                "success": False,
                "error": "no OPENAI_API_KEY",
                "context_files": sources,
            }
        return {"success": False, "error": llm.get("error"), "context_files": sources}

    return {"success": True, "answer": llm.get("content"), "sources": sources}


def _file_evaluate(
    file_path: str,
    expected_intent: str,
    source_data_description=None,
    criteria=None,
    model: str = "gpt-4o-mini",
) -> dict:
    """LLM quality check of a produced file against an expected intent.

    Reads the file as text; for .xlsx/.docx extracts text via openpyxl/
    python-docx when available, else notes it's binary. Returns a verdict
    (pass|fail|partial), a 0..1 score, and reasons. Graceful no-key fallback.
    """
    if not os.path.isfile(file_path):
        return {"success": False, "error": f"file not found: {file_path}"}

    ext = os.path.splitext(file_path)[1].lower()
    content = ""
    try:
        if ext == ".xlsx":
            try:
                import openpyxl  # type: ignore

                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                rows_out = []
                for ws in wb.worksheets:
                    rows_out.append(f"# Sheet: {ws.title}")
                    for i, row in enumerate(ws.iter_rows(values_only=True)):
                        if i >= 200:
                            rows_out.append("... (truncated)")
                            break
                        rows_out.append(
                            "\t".join("" if c is None else str(c) for c in row)
                        )
                content = "\n".join(rows_out)
                wb.close()
            except Exception as e:
                content = f"[could not extract .xlsx text: {e}]"
        elif ext == ".docx":
            try:
                import docx  # type: ignore

                doc = docx.Document(file_path)
                content = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                content = f"[could not extract .docx text: {e}]"
        else:
            with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
                content = fh.read()
    except Exception as e:
        return {"success": False, "error": f"could not read file: {e}"}

    content_for_llm = content[:12000]
    crit_text = ""
    if criteria:
        crit_text = "\nCRITERIA:\n" + (
            "\n".join(f"- {c}" for c in criteria)
            if isinstance(criteria, (list, tuple))
            else str(criteria)
        )
    src_text = (
        f"\nSOURCE DATA DESCRIPTION:\n{source_data_description}"
        if source_data_description
        else ""
    )

    prompt = (
        "You are a strict QA reviewer. Judge whether the FILE CONTENT satisfies "
        "the EXPECTED INTENT (and any criteria). Respond ONLY with a JSON object: "
        '{"verdict":"pass|fail|partial","score":0.0-1.0,"reasons":"..."}\n\n'
        f"EXPECTED INTENT:\n{expected_intent}{crit_text}{src_text}\n\n"
        f"FILE: {file_path}\nFILE CONTENT:\n{content_for_llm}"
    )
    llm = _llm_chat_complete(
        [
            {"role": "system", "content": "You output only valid JSON."},
            {"role": "user", "content": prompt},
        ],
        model=model,
    )
    if not llm.get("success"):
        if llm.get("error") == "no OPENAI_API_KEY":
            return {"success": False, "error": "no OPENAI_API_KEY", "file": file_path}
        return {"success": False, "error": llm.get("error"), "file": file_path}

    raw = (llm.get("content") or "").strip()
    verdict, score, reasons = "partial", 0.5, raw
    try:
        text = raw
        if text.startswith("```"):
            text = text.strip("`")
            if text.lower().startswith("json"):
                text = text[4:]
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end != -1:
            parsed = json.loads(text[start : end + 1])
            verdict = str(parsed.get("verdict", verdict)).lower()
            score = float(parsed.get("score", score))
            reasons = parsed.get("reasons", reasons)
    except Exception:
        pass
    if verdict not in ("pass", "fail", "partial"):
        verdict = "partial"
    score = max(0.0, min(1.0, score))

    return {
        "success": True,
        "verdict": verdict,
        "score": score,
        "reasons": reasons,
        "file": file_path,
    }


# ─── Tool() schemas for the 8 data / knowledge handlers above ───
TOOLS += [
    Tool(
        name="rowboat_upload",
        description="Upload a knowledge document (file) to the local Rowboat knowledge base. NOTE: the upload endpoint is not verified for the running Rowboat version and may differ; failures return a clean error with a hint.",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path to the file to upload.",
                },
                "title": {
                    "type": "string",
                    "description": "Optional document title.",
                },
                "tags": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional list of tags.",
                },
            },
            "required": ["file_path"],
        },
    ),
    Tool(
        name="rowboat_chat",
        description="Chat with the local Rowboat agent (semantic KB-backed). Returns the agent's answer plus the conversation id for follow-ups.",
        inputSchema={
            "type": "object",
            "properties": {
                "message": {
                    "type": "string",
                    "description": "The user message / question to send.",
                },
                "context": {
                    "type": "string",
                    "default": "default",
                    "description": "Optional context label passed to Rowboat.",
                },
                "conversation_id": {
                    "type": "string",
                    "description": "Optional existing conversation id to continue a thread.",
                },
                "timeout": {
                    "type": "integer",
                    "default": 120,
                    "description": "HTTP timeout in seconds.",
                },
            },
            "required": ["message"],
        },
    ),
    Tool(
        name="rowboat_search",
        description="Semantic search over the Rowboat knowledge base. NOTE: the search endpoint is not verified for the running Rowboat version; failures return a clean error with a hint.",
        inputSchema={
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query.",
                },
                "folder": {
                    "type": "string",
                    "description": "Optional folder/source to restrict the search to.",
                },
                "limit": {
                    "type": "integer",
                    "default": 20,
                    "description": "Maximum number of results.",
                },
            },
            "required": ["query"],
        },
    ),
    Tool(
        name="rowboat_list_folders",
        description="List Rowboat knowledge folders/sources. NOTE: the sources endpoint is not verified for the running Rowboat version; failures return a clean error with a hint.",
        inputSchema={
            "type": "object",
            "properties": {
                "limit": {
                    "type": "integer",
                    "default": 50,
                    "description": "Maximum number of sources to return.",
                },
            },
        },
    ),
    Tool(
        name="roarboot_ask",
        description="RAG question-answering over a local markdown knowledge vault (filesystem, NOT Rowboat). Ranks .md files in a folder by overlap with the question, builds a context prompt, and asks an LLM. Requires OPENAI_API_KEY for the answer step (otherwise reports the files it would have used).",
        inputSchema={
            "type": "object",
            "properties": {
                "question": {
                    "type": "string",
                    "description": "The question to answer from the vault.",
                },
                "folder": {
                    "type": "string",
                    "description": "Optional subfolder under the knowledge root to restrict to.",
                },
                "max_files": {
                    "type": "integer",
                    "default": 10,
                    "description": "Maximum number of files to include as context.",
                },
                "max_chars_per_file": {
                    "type": "integer",
                    "default": 4000,
                    "description": "Truncate each file's content to this many characters.",
                },
                "model": {
                    "type": "string",
                    "default": "gpt-4o-mini",
                    "description": "LLM model name for the answer step.",
                },
                "root": {
                    "type": "string",
                    "description": "Optional override of the knowledge root directory.",
                },
            },
            "required": ["question"],
        },
    ),
    Tool(
        name="roarboot_read_knowledge",
        description="Read .md files in <root>/<folder> (recursively) from a local markdown knowledge vault. Optionally filter filenames by name_pattern (glob or substring) and/or content by query (case-insensitive substring). Returns file paths plus short excerpts.",
        inputSchema={
            "type": "object",
            "properties": {
                "folder": {
                    "type": "string",
                    "description": "Subfolder under the knowledge root to read.",
                },
                "query": {
                    "type": "string",
                    "description": "Optional case-insensitive substring to filter file content.",
                },
                "limit": {
                    "type": "integer",
                    "default": 20,
                    "description": "Maximum number of files to return.",
                },
                "name_pattern": {
                    "type": "string",
                    "description": "Optional glob (e.g. '*plan*.md') or substring to filter filenames.",
                },
                "root": {
                    "type": "string",
                    "description": "Optional override of the knowledge root directory.",
                },
            },
            "required": ["folder"],
        },
    ),
    Tool(
        name="roarboot_list_folders",
        description="List the immediate subfolders under the local markdown knowledge vault root (filesystem, NOT Rowboat).",
        inputSchema={
            "type": "object",
            "properties": {
                "root": {
                    "type": "string",
                    "description": "Optional override of the knowledge root directory.",
                },
            },
        },
    ),
    Tool(
        name="file_evaluate",
        description="LLM quality check of a produced file: reads the file (extracts text from .xlsx/.docx when possible) and judges whether it matches an expected intent (plus optional criteria). Returns verdict (pass|fail|partial), a 0..1 score, and reasons. Requires OPENAI_API_KEY.",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path to the file to evaluate.",
                },
                "expected_intent": {
                    "type": "string",
                    "description": "What the file was supposed to accomplish/contain.",
                },
                "source_data_description": {
                    "type": "string",
                    "description": "Optional description of the source data the file was built from.",
                },
                "criteria": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional explicit acceptance criteria.",
                },
                "model": {
                    "type": "string",
                    "default": "gpt-4o-mini",
                    "description": "LLM model name for the evaluation.",
                },
            },
            "required": ["file_path", "expected_intent"],
        },
    ),
]


# ─── Phase: deterministic document creation (4-way-split document server) ───
#
# Sync handlers consumed by the "document" MCP entrypoint (imports this module
# as H and calls these by exact name/signature, filtering H.TOOLS by name).
# All imports of openpyxl / python-docx are lazy (mirrors the file's style).


def _xlsx_create_from_data(
    file_path: str,
    rows=None,
    sheet_name=None,
    bold_rows=None,
    auto_width=True,
    overwrite=True,
    sheets=None,
    cell_styles=None,
    freeze_pane=None,
) -> dict:
    """Deterministic .xlsx creation via openpyxl.

    rows = list of row-lists (first may be header). If `sheets` (dict
    sheet_name->rows) is given, create multiple sheets; else one sheet
    named `sheet_name or "Sheet1"` from `rows`. `bold_rows` = 0-based row
    indices to bold. `auto_width` sizes columns to content. `freeze_pane`
    e.g. "A2". `cell_styles` best-effort (unknown keys ignored).
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
        from openpyxl.utils import get_column_letter

        abs_path = os.path.abspath(file_path)
        if not overwrite and os.path.exists(abs_path):
            return {"success": False, "error": "exists", "file_path": abs_path}

        parent = os.path.dirname(abs_path)
        if parent:
            os.makedirs(parent, exist_ok=True)

        # Normalize into an ordered list of (name, rows) sheet specs.
        sheet_specs = []
        if sheets and isinstance(sheets, dict):
            for s_name, s_rows in sheets.items():
                sheet_specs.append((str(s_name), list(s_rows or [])))
        else:
            sheet_specs.append((str(sheet_name or "Sheet1"), list(rows or [])))

        bold_set = set(int(i) for i in (bold_rows or []))

        wb = Workbook()
        # Remove the default sheet; we add our own named ones.
        default_ws = wb.active
        wb.remove(default_ws)

        total_rows = 0
        sheet_names = []
        for s_name, s_rows in sheet_specs:
            # Excel sheet names max 31 chars and cannot contain some chars.
            safe_name = str(s_name)[:31] or "Sheet"
            ws = wb.create_sheet(title=safe_name)
            sheet_names.append(ws.title)

            max_widths = {}
            for r_idx, row in enumerate(s_rows):
                cells = list(row) if isinstance(row, (list, tuple)) else [row]
                for c_idx, value in enumerate(cells):
                    cell = ws.cell(row=r_idx + 1, column=c_idx + 1, value=value)
                    if r_idx in bold_set:
                        cell.font = Font(bold=True)
                    if auto_width:
                        text = "" if value is None else str(value)
                        col = c_idx + 1
                        max_widths[col] = max(max_widths.get(col, 0), len(text))
                total_rows += 1

            if auto_width:
                for col, width in max_widths.items():
                    ws.column_dimensions[get_column_letter(col)].width = min(
                        max(width + 2, 8), 80
                    )

            if freeze_pane:
                try:
                    ws.freeze_panes = str(freeze_pane)
                except Exception:
                    pass

            # Best-effort cell_styles: {"A1": {"bold": True, "italic": True,
            # "size": int}}. Only applies to the first/only sheet unless a
            # dict-of-dicts keyed by sheet name is provided.
            if cell_styles and isinstance(cell_styles, dict):
                styles_for_sheet = cell_styles
                if safe_name in cell_styles and isinstance(
                    cell_styles[safe_name], dict
                ):
                    styles_for_sheet = cell_styles[safe_name]
                for ref, spec in styles_for_sheet.items():
                    if not isinstance(spec, dict):
                        continue
                    try:
                        cell = ws[ref]
                    except Exception:
                        continue
                    try:
                        cell.font = Font(
                            bold=bool(spec.get("bold", False)),
                            italic=bool(spec.get("italic", False)),
                            size=int(spec["size"]) if spec.get("size") else None,
                        )
                    except Exception:
                        pass

        wb.save(abs_path)
        return {
            "success": True,
            "file_path": abs_path,
            "sheets": sheet_names,
            "rows_written": total_rows,
        }
    except Exception as e:
        logger.error(f"_xlsx_create_from_data failed: {e}")
        return {"success": False, "error": str(e), "file_path": file_path}


def _docx_create_from_data(file_path: str, blocks=None, overwrite=True) -> dict:
    """Create a .docx via python-docx from a list of block dicts.

    block types: heading {"text","level"}, paragraph {"text"},
    bullet {"items":[...]}, table {"rows":[[...]],"header":bool}.
    """
    try:
        from docx import Document

        abs_path = os.path.abspath(file_path)
        if not overwrite and os.path.exists(abs_path):
            return {"success": False, "error": "exists", "file_path": abs_path}

        parent = os.path.dirname(abs_path)
        if parent:
            os.makedirs(parent, exist_ok=True)

        doc = Document()
        block_list = list(blocks or [])
        for block in block_list:
            if not isinstance(block, dict):
                continue
            btype = str(block.get("type", "paragraph")).lower()
            if btype == "heading":
                level = int(block.get("level", 1) or 1)
                doc.add_heading(str(block.get("text", "")), level=max(0, min(level, 9)))
            elif btype == "paragraph":
                doc.add_paragraph(str(block.get("text", "")))
            elif btype == "bullet":
                for item in block.get("items", []) or []:
                    doc.add_paragraph(str(item), style="List Bullet")
            elif btype == "table":
                t_rows = block.get("rows", []) or []
                if not t_rows:
                    continue
                ncols = max(len(r) for r in t_rows)
                table = doc.add_table(rows=0, cols=ncols)
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
                has_header = bool(block.get("header", False))
                for r_idx, row in enumerate(t_rows):
                    cells = list(row) if isinstance(row, (list, tuple)) else [row]
                    tr = table.add_row().cells
                    for c_idx in range(ncols):
                        val = cells[c_idx] if c_idx < len(cells) else ""
                        tr[c_idx].text = "" if val is None else str(val)
                        if has_header and r_idx == 0:
                            for para in tr[c_idx].paragraphs:
                                for run in para.runs:
                                    run.bold = True
            else:
                # Unknown block type: best-effort treat as paragraph text.
                if "text" in block:
                    doc.add_paragraph(str(block.get("text", "")))

        doc.save(abs_path)
        return {"success": True, "file_path": abs_path, "blocks": len(block_list)}
    except Exception as e:
        logger.error(f"_docx_create_from_data failed: {e}")
        return {"success": False, "error": str(e), "file_path": file_path}


def _csv_create_from_data(
    file_path: str,
    rows=None,
    delimiter=",",
    encoding="utf-8-sig",
    overwrite=True,
) -> dict:
    """Write rows (list of lists) to CSV via the stdlib csv module."""
    try:
        import csv as _csv

        abs_path = os.path.abspath(file_path)
        if not overwrite and os.path.exists(abs_path):
            return {"success": False, "error": "exists", "file_path": abs_path}

        parent = os.path.dirname(abs_path)
        if parent:
            os.makedirs(parent, exist_ok=True)

        row_list = list(rows or [])
        with open(abs_path, "w", newline="", encoding=encoding) as fh:
            writer = _csv.writer(fh, delimiter=delimiter)
            for row in row_list:
                if isinstance(row, (list, tuple)):
                    writer.writerow(list(row))
                else:
                    writer.writerow([row])

        return {"success": True, "file_path": abs_path, "rows": len(row_list)}
    except Exception as e:
        logger.error(f"_csv_create_from_data failed: {e}")
        return {"success": False, "error": str(e), "file_path": file_path}


def _excel_verify_file(
    file_path: str,
    expected_cells=None,
    min_rows=None,
    must_contain_text=None,
) -> dict:
    """Open an .xlsx (read-only) and verify cells/row-count/text presence."""
    checks = []
    failures = []
    try:
        from openpyxl import load_workbook

        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            return {
                "success": True,
                "valid": False,
                "checks": [],
                "failures": [f"file not found: {abs_path}"],
            }

        wb = load_workbook(abs_path, read_only=True, data_only=True)
        ws = wb.active

        if expected_cells and isinstance(expected_cells, dict):
            for ref, expected in expected_cells.items():
                try:
                    actual = ws[ref].value
                except Exception:
                    actual = None
                ok = ("" if actual is None else str(actual)) == str(expected)
                msg = f"cell {ref} == {expected!r} (got {actual!r})"
                checks.append({"check": msg, "ok": ok})
                if not ok:
                    failures.append(msg)

        if min_rows is not None:
            actual_rows = ws.max_row or 0
            ok = actual_rows >= int(min_rows)
            msg = f"min_rows >= {min_rows} (got {actual_rows})"
            checks.append({"check": msg, "ok": ok})
            if not ok:
                failures.append(msg)

        if must_contain_text:
            # Collect all stringified cell values once.
            haystack_parts = []
            for row in ws.iter_rows(values_only=True):
                for val in row:
                    if val is not None:
                        haystack_parts.append(str(val))
            haystack = "\n".join(haystack_parts)
            for needle in must_contain_text:
                ok = str(needle) in haystack
                msg = f"contains text {needle!r}"
                checks.append({"check": msg, "ok": ok})
                if not ok:
                    failures.append(msg)

        try:
            wb.close()
        except Exception:
            pass

        return {
            "success": True,
            "valid": len(failures) == 0,
            "checks": checks,
            "failures": failures,
        }
    except Exception as e:
        logger.error(f"_excel_verify_file failed: {e}")
        return {
            "success": False,
            "valid": False,
            "checks": checks,
            "failures": failures + [str(e)],
            "error": str(e),
        }


def _docx_verify_file(
    file_path: str,
    expected_substrings=None,
    min_paragraphs=None,
    min_tables=None,
    expected_table_cells=None,
) -> dict:
    """Open a .docx and verify substrings / paragraph & table counts / cells."""
    checks = []
    failures = []
    try:
        from docx import Document

        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            return {
                "success": True,
                "valid": False,
                "checks": [],
                "failures": [f"file not found: {abs_path}"],
            }

        doc = Document(abs_path)
        paragraphs = doc.paragraphs
        tables = doc.tables
        full_text = "\n".join(p.text for p in paragraphs)
        # Include table cell text in the searchable full text.
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        if expected_substrings:
            for needle in expected_substrings:
                ok = str(needle) in full_text
                msg = f"contains substring {needle!r}"
                checks.append({"check": msg, "ok": ok})
                if not ok:
                    failures.append(msg)

        if min_paragraphs is not None:
            ok = len(paragraphs) >= int(min_paragraphs)
            msg = f"min_paragraphs >= {min_paragraphs} (got {len(paragraphs)})"
            checks.append({"check": msg, "ok": ok})
            if not ok:
                failures.append(msg)

        if min_tables is not None:
            ok = len(tables) >= int(min_tables)
            msg = f"min_tables >= {min_tables} (got {len(tables)})"
            checks.append({"check": msg, "ok": ok})
            if not ok:
                failures.append(msg)

        if expected_table_cells:
            for spec in expected_table_cells:
                if not isinstance(spec, dict):
                    continue
                t_i = int(spec.get("table", 0))
                r_i = int(spec.get("row", 0))
                c_i = int(spec.get("col", 0))
                expected = str(spec.get("text", ""))
                actual = None
                try:
                    actual = tables[t_i].rows[r_i].cells[c_i].text
                except Exception:
                    actual = None
                ok = actual == expected
                msg = (
                    f"table[{t_i}][{r_i}][{c_i}] == {expected!r} "
                    f"(got {actual!r})"
                )
                checks.append({"check": msg, "ok": ok})
                if not ok:
                    failures.append(msg)

        return {
            "success": True,
            "valid": len(failures) == 0,
            "checks": checks,
            "failures": failures,
        }
    except Exception as e:
        logger.error(f"_docx_verify_file failed: {e}")
        return {
            "success": False,
            "valid": False,
            "checks": checks,
            "failures": failures + [str(e)],
            "error": str(e),
        }


def _excel_paste_table(rows=None, start_cell=None) -> dict:
    """Paste a table into the currently active Excel workbook via COM.

    Reuses agents.handoff.office_automation._get_excel to obtain the running
    Excel instance / active workbook. Never raises; returns a clean error
    when no Excel instance / COM is available.
    """
    try:
        from agents.handoff.office_automation import _get_excel

        start = str(start_cell or "A1")
        row_list = list(rows or [])

        xl = _get_excel()
        if xl is None:
            return {"success": False, "error": "no active Excel workbook"}

        try:
            wb = xl.ActiveWorkbook
        except Exception:
            wb = None
        if wb is None:
            return {"success": False, "error": "no active Excel workbook"}

        try:
            ws = wb.ActiveSheet
        except Exception:
            return {"success": False, "error": "no active Excel workbook"}

        # Resolve the top-left anchor cell, then write row-by-row by offset.
        try:
            anchor = ws.Range(start)
        except Exception:
            return {"success": False, "error": f"invalid start_cell: {start}"}

        base_row = anchor.Row
        base_col = anchor.Column
        cells_written = 0
        for r_idx, row in enumerate(row_list):
            cells = list(row) if isinstance(row, (list, tuple)) else [row]
            for c_idx, value in enumerate(cells):
                try:
                    ws.Cells(base_row + r_idx, base_col + c_idx).Value = value
                    cells_written += 1
                except Exception:
                    pass

        return {
            "success": True,
            "cells_written": cells_written,
            "start_cell": start,
        }
    except Exception as e:
        logger.error(f"_excel_paste_table failed: {e}")
        return {"success": False, "error": "no active Excel workbook", "detail": str(e)}


def _window_maximize(title_substring: str) -> dict:
    """Maximize the first visible window whose title contains the substring."""
    try:
        import ctypes

        from agents.handoff.window_focus import find_window_by_title

        hwnd = find_window_by_title(title_substring)
        if not hwnd:
            return {"success": False, "error": "window not found"}

        # SW_MAXIMIZE = 3
        ctypes.windll.user32.ShowWindow(hwnd, 3)
        return {"success": True, "hwnd": int(hwnd)}
    except Exception as e:
        logger.error(f"_window_maximize failed: {e}")
        return {"success": False, "error": str(e)}


TOOLS += [
    Tool(
        name="xlsx_create_from_data",
        description="Deterministically create an .xlsx file from row data via openpyxl (no LLM). Supports a single sheet from `rows` or multiple sheets via `sheets`, header bolding, auto column width, and freeze panes.",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path of the .xlsx file to create.",
                },
                "rows": {
                    "type": "array",
                    "items": {"type": "array"},
                    "description": "List of row-lists for the (single) sheet. First row may be a header.",
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name for the single sheet (default 'Sheet1'). Ignored if `sheets` is given.",
                },
                "bold_rows": {
                    "type": "array",
                    "items": {"type": "integer"},
                    "description": "0-based row indices to render bold (e.g. [0] for a header).",
                },
                "auto_width": {
                    "type": "boolean",
                    "default": True,
                    "description": "Size columns to fit their content.",
                },
                "overwrite": {
                    "type": "boolean",
                    "default": True,
                    "description": "If False and the file exists, returns an 'exists' error.",
                },
                "sheets": {
                    "type": "object",
                    "description": "Mapping sheet_name -> rows (list of row-lists) to create multiple sheets.",
                },
                "cell_styles": {
                    "type": "object",
                    "description": "Optional best-effort per-cell styles, e.g. {'A1': {'bold': true, 'size': 14}}.",
                },
                "freeze_pane": {
                    "type": "string",
                    "description": "Optional freeze-pane anchor, e.g. 'A2'.",
                },
            },
            "required": ["file_path"],
        },
    ),
    Tool(
        name="docx_create_from_data",
        description="Deterministically create a .docx file from structured blocks via python-docx (no LLM). Block types: heading, paragraph, bullet, table.",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path of the .docx file to create.",
                },
                "blocks": {
                    "type": "array",
                    "items": {"type": "object"},
                    "description": "List of block dicts: heading {text,level}, paragraph {text}, bullet {items:[...]}, table {rows:[[...]],header:bool}.",
                },
                "overwrite": {
                    "type": "boolean",
                    "default": True,
                    "description": "If False and the file exists, returns an 'exists' error.",
                },
            },
            "required": ["file_path"],
        },
    ),
    Tool(
        name="csv_create_from_data",
        description="Deterministically write rows (list of lists) to a CSV file via the stdlib csv module.",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path of the .csv file to create.",
                },
                "rows": {
                    "type": "array",
                    "items": {"type": "array"},
                    "description": "List of row-lists to write.",
                },
                "delimiter": {
                    "type": "string",
                    "default": ",",
                    "description": "Field delimiter.",
                },
                "encoding": {
                    "type": "string",
                    "default": "utf-8-sig",
                    "description": "File encoding (utf-8-sig keeps Excel happy with BOM).",
                },
                "overwrite": {
                    "type": "boolean",
                    "default": True,
                    "description": "If False and the file exists, returns an 'exists' error.",
                },
            },
            "required": ["file_path"],
        },
    ),
    Tool(
        name="excel_verify_file",
        description="Open an existing .xlsx (read-only) and verify it: expected cell values, a minimum active-sheet row count, and/or required text. Returns valid plus per-check results and failures.",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path of the .xlsx file to verify.",
                },
                "expected_cells": {
                    "type": "object",
                    "description": "Mapping cell-ref -> expected value, e.g. {'A1': 'Name'} (compared as strings).",
                },
                "min_rows": {
                    "type": "integer",
                    "description": "Active sheet must have at least this many rows.",
                },
                "must_contain_text": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Strings that must appear somewhere in the sheet.",
                },
            },
            "required": ["file_path"],
        },
    ),
    Tool(
        name="docx_verify_file",
        description="Open an existing .docx and verify it: required substrings, minimum paragraph/table counts, and/or specific table cell values. Returns valid plus per-check results and failures.",
        inputSchema={
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path of the .docx file to verify.",
                },
                "expected_substrings": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Substrings that must be present in the document text.",
                },
                "min_paragraphs": {
                    "type": "integer",
                    "description": "Document must have at least this many paragraphs.",
                },
                "min_tables": {
                    "type": "integer",
                    "description": "Document must have at least this many tables.",
                },
                "expected_table_cells": {
                    "type": "array",
                    "items": {"type": "object"},
                    "description": "List of {table,row,col,text} expectations (0-based indices).",
                },
            },
            "required": ["file_path"],
        },
    ),
    Tool(
        name="excel_paste_table",
        description="Paste a table of rows into the CURRENTLY ACTIVE Excel workbook via COM (reuses office_automation). Returns a clean error if no Excel instance is running.",
        inputSchema={
            "type": "object",
            "properties": {
                "rows": {
                    "type": "array",
                    "items": {"type": "array"},
                    "description": "List of row-lists to write into the active sheet.",
                },
                "start_cell": {
                    "type": "string",
                    "default": "A1",
                    "description": "Top-left anchor cell (default 'A1').",
                },
            },
            "required": ["rows"],
        },
    ),
    Tool(
        name="window_maximize",
        description="Maximize the first visible window whose title contains the given substring (reuses window_focus.find_window_by_title + SW_MAXIMIZE).",
        inputSchema={
            "type": "object",
            "properties": {
                "title_substring": {
                    "type": "string",
                    "description": "Substring to match against window titles.",
                },
            },
            "required": ["title_substring"],
        },
    ),
]


if __name__ == "__main__":
    asyncio.run(main())
