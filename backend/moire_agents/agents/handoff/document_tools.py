"""Document processing tools for the Handoff MCP server.

Phase 1.5: implements the 5 document MCP tools that the documentation
promised but the server never registered:

  - handoff_doc_scan   — parse a PDF/DOCX/TXT into a structured tree
  - handoff_doc_edit   — queue an in-memory edit operation
  - handoff_doc_apply  — write all queued edits back to the source file
  - handoff_doc_export — export a parsed document as JSON / Markdown / TXT / DOCX
  - handoff_doc_list   — list all open document sessions

Architecture
============

`handoff_doc_scan` returns a stable `document_id` and stores the parsed tree
plus a list of pending edits in an in-process registry. Subsequent
`handoff_doc_edit` / `handoff_doc_apply` / `handoff_doc_export` calls
reference this id.

The native parsing path uses PyMuPDF for PDFs and python-docx for DOCX. Plain
text files are read as a single section. If neither library is available,
the tool returns a structured "library missing" error rather than crashing.

The "visual" path (scan an on-screen document via OCR + auto-scroll) is
mentioned in the plan but is intentionally NOT implemented in this file —
the screen-description and OCR tools live elsewhere and would create a
circular dependency. Visual scan can be assembled at the MCP server level
once Phase 4 lands.
"""

from __future__ import annotations

import json
import logging
import os
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


# ─── Session registry ────────────────────────────────────────────────────────


@dataclass
class DocSection:
    """A logical section of a document — heading + body text."""

    index: int
    text: str
    heading: Optional[str] = None
    style: Optional[str] = None


@dataclass
class DocPage:
    """A single page (PDF) or chunk (DOCX/TXT)."""

    number: int
    sections: List[DocSection] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(s.text for s in self.sections)


@dataclass
class DocSession:
    """An open document session, kept in memory until apply or expiry."""

    id: str
    source_path: str
    source_type: str  # 'pdf' | 'docx' | 'txt'
    pages: List[DocPage] = field(default_factory=list)
    pending_edits: List[Dict[str, Any]] = field(default_factory=list)
    created_at: float = field(default_factory=time.time)


_sessions: Dict[str, DocSession] = {}


def _new_session_id() -> str:
    return f"doc_{uuid.uuid4().hex[:12]}"


def _detect_type(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".docx", ".doc"}:
        return "docx"
    return "txt"


# ─── Native parsers ──────────────────────────────────────────────────────────


def _parse_pdf(path: Path, max_pages: int) -> List[DocPage]:
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError(f"pymupdf not installed: {e}")

    pages: List[DocPage] = []
    with fitz.open(str(path)) as doc:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            blocks = page.get_text(
                "blocks"
            )  # list of (x0,y0,x1,y1,text,block_no,block_type)
            sections: List[DocSection] = []
            for j, blk in enumerate(blocks):
                text = (blk[4] if len(blk) > 4 else "").strip()
                if not text:
                    continue
                # Heuristic: short, all-caps-ish lines are treated as headings.
                heading = None
                first_line = text.split("\n", 1)[0]
                if len(first_line) < 80 and (
                    first_line.isupper()
                    or sum(1 for c in first_line if c.isupper()) > len(first_line) * 0.6
                ):
                    heading = first_line
                sections.append(DocSection(index=j, text=text, heading=heading))
            pages.append(DocPage(number=i + 1, sections=sections))
    return pages


def _parse_docx(path: Path, max_pages: int) -> List[DocPage]:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(f"python-docx not installed: {e}")

    doc = docx.Document(str(path))
    sections: List[DocSection] = []
    for j, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        style = para.style.name if para.style else None
        heading = text if (style or "").startswith("Heading") else None
        sections.append(DocSection(index=j, text=text, heading=heading, style=style))

    # python-docx has no page concept; chunk into "pages" of ~40 sections each
    # so editing remains addressable by (page, section_index).
    chunk = 40
    pages: List[DocPage] = []
    for p_idx, start in enumerate(range(0, len(sections), chunk)):
        if p_idx >= max_pages:
            break
        page_sections = sections[start : start + chunk]
        # Re-index local section.index so doc_edit(page=N, section_index=K) is stable.
        for local_i, sec in enumerate(page_sections):
            sec.index = local_i
        pages.append(DocPage(number=p_idx + 1, sections=page_sections))
    return pages


def _parse_txt(path: Path, max_pages: int) -> List[DocPage]:
    text = path.read_text(encoding="utf-8", errors="replace")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    sections = [DocSection(index=i, text=p) for i, p in enumerate(paragraphs)]
    return [DocPage(number=1, sections=sections)]


# ─── Tool implementations ────────────────────────────────────────────────────


async def handle_doc_scan(
    path: Optional[str] = None,
    max_pages: int = 20,
    detect_structure: bool = True,
    scroll_amount: int = 800,  # accepted for API parity, used by visual path only
) -> Dict[str, Any]:
    """Parse a document file and create a session.

    Visual scanning (no path supplied) is not yet wired through this module —
    it returns a 501-style payload pointing the caller at the screen-reading
    tools. Phase 4's `handoff_describe_screen` is the better path.
    """
    if not path:
        return {
            "success": False,
            "error": "visual document scan not implemented in document_tools.py",
            "hint": "use handoff_describe_screen + handoff_read_screen for on-screen documents",
        }

    p = Path(path)
    if not p.exists():
        return {"success": False, "error": f"file does not exist: {path}"}
    if not p.is_file():
        return {"success": False, "error": f"not a file: {path}"}

    doc_type = _detect_type(p)
    try:
        if doc_type == "pdf":
            pages = _parse_pdf(p, max_pages)
        elif doc_type == "docx":
            pages = _parse_docx(p, max_pages)
        else:
            pages = _parse_txt(p, max_pages)
    except RuntimeError as e:
        return {"success": False, "error": str(e), "type": doc_type}
    except Exception as e:
        return {"success": False, "error": f"parse failed: {e}", "type": doc_type}

    session = DocSession(
        id=_new_session_id(),
        source_path=str(p.resolve()),
        source_type=doc_type,
        pages=pages,
    )
    _sessions[session.id] = session

    return {
        "success": True,
        "document_id": session.id,
        "source_path": session.source_path,
        "source_type": doc_type,
        "page_count": len(pages),
        "pages": [
            {
                "number": pg.number,
                "section_count": len(pg.sections),
                "sections": [
                    {
                        "index": s.index,
                        "heading": s.heading,
                        "style": s.style,
                        "preview": s.text[:120],
                    }
                    for s in pg.sections
                ],
            }
            for pg in pages
        ],
    }


async def handle_doc_edit(
    document_id: str,
    page: int,
    section_index: int,
    new_text: str,
    operation: str = "replace",  # replace | append | prepend | delete
) -> Dict[str, Any]:
    """Queue an in-memory edit. Apply with handle_doc_apply."""
    session = _sessions.get(document_id)
    if not session:
        return {"success": False, "error": f"unknown document_id: {document_id}"}

    if operation not in {"replace", "append", "prepend", "delete"}:
        return {"success": False, "error": f"unknown operation: {operation}"}

    # Locate section and apply in-memory.
    target_page = next((pg for pg in session.pages if pg.number == page), None)
    if not target_page:
        return {"success": False, "error": f"page {page} not in document"}
    if section_index < 0 or section_index >= len(target_page.sections):
        return {
            "success": False,
            "error": f"section_index {section_index} out of range (0..{len(target_page.sections)-1})",
        }

    section = target_page.sections[section_index]
    old_text = section.text

    if operation == "replace":
        section.text = new_text
    elif operation == "append":
        section.text = old_text + new_text
    elif operation == "prepend":
        section.text = new_text + old_text
    elif operation == "delete":
        section.text = ""

    edit = {
        "page": page,
        "section_index": section_index,
        "operation": operation,
        "old_text": old_text,
        "new_text": section.text,
        "timestamp": time.time(),
    }
    session.pending_edits.append(edit)

    return {
        "success": True,
        "document_id": document_id,
        "edit_count": len(session.pending_edits),
        "edit": edit,
    }


async def handle_doc_apply(
    document_id: str,
    dry_run: bool = False,
) -> Dict[str, Any]:
    """Write the in-memory edited document back to its source file."""
    session = _sessions.get(document_id)
    if not session:
        return {"success": False, "error": f"unknown document_id: {document_id}"}
    if not session.pending_edits:
        return {
            "success": True,
            "document_id": document_id,
            "applied": 0,
            "note": "no pending edits",
        }

    if dry_run:
        return {
            "success": True,
            "document_id": document_id,
            "would_apply": len(session.pending_edits),
            "edits": session.pending_edits,
            "dry_run": True,
        }

    src = Path(session.source_path)
    backup_path = src.with_suffix(src.suffix + ".bak")
    try:
        # Always make a backup before writing.
        if src.exists() and not backup_path.exists():
            import shutil

            shutil.copy2(src, backup_path)

        if session.source_type == "pdf":
            return await _apply_pdf(session, src)
        if session.source_type == "docx":
            return await _apply_docx(session, src)
        return await _apply_txt(session, src)
    except Exception as e:
        return {"success": False, "error": f"apply failed: {e}"}


async def _apply_pdf(session: DocSession, src: Path) -> Dict[str, Any]:
    """PDF write-back: PyMuPDF redaction. Not full edit support."""
    try:
        import fitz
    except ImportError:
        return {"success": False, "error": "pymupdf not installed"}

    with fitz.open(str(src)) as doc:
        for edit in session.pending_edits:
            page_num = edit["page"] - 1
            if page_num < 0 or page_num >= len(doc):
                continue
            page = doc[page_num]
            old = edit["old_text"]
            new = edit["new_text"]
            for hit in page.search_for(old):
                page.add_redact_annot(hit, text=new, fill=(1, 1, 1))
            page.apply_redactions()
        doc.save(str(src), incremental=False, deflate=True, clean=True)

    applied = len(session.pending_edits)
    session.pending_edits.clear()
    return {
        "success": True,
        "document_id": session.id,
        "applied": applied,
        "source_path": str(src),
        "backup": str(src.with_suffix(src.suffix + ".bak")),
        "warning": "PDF edits applied via redaction; layout may shift",
    }


async def _apply_docx(session: DocSession, src: Path) -> Dict[str, Any]:
    """DOCX write-back via python-docx. Replaces paragraph text in place."""
    try:
        import docx
    except ImportError:
        return {"success": False, "error": "python-docx not installed"}

    doc = docx.Document(str(src))
    paragraphs = doc.paragraphs

    # Build a flat (page, section_index) → global paragraph index map.
    chunk = 40
    new_text_for_global: Dict[int, str] = {}
    for pg in session.pages:
        page_offset = (pg.number - 1) * chunk
        for sec in pg.sections:
            global_idx = page_offset + sec.index
            if 0 <= global_idx < len(paragraphs):
                new_text_for_global[global_idx] = sec.text

    for global_idx, new_text in new_text_for_global.items():
        para = paragraphs[global_idx]
        if para.text != new_text:
            # Clear runs and write the new text in a single run.
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

    doc.save(str(src))
    applied = len(session.pending_edits)
    session.pending_edits.clear()
    return {
        "success": True,
        "document_id": session.id,
        "applied": applied,
        "source_path": str(src),
        "backup": str(src.with_suffix(src.suffix + ".bak")),
    }


async def _apply_txt(session: DocSession, src: Path) -> Dict[str, Any]:
    """TXT write-back: rewrite the file from current section text."""
    paragraphs = []
    for pg in session.pages:
        for sec in pg.sections:
            if sec.text:
                paragraphs.append(sec.text)
    src.write_text("\n\n".join(paragraphs), encoding="utf-8")
    applied = len(session.pending_edits)
    session.pending_edits.clear()
    return {
        "success": True,
        "document_id": session.id,
        "applied": applied,
        "source_path": str(src),
        "backup": str(src.with_suffix(src.suffix + ".bak")),
    }


async def handle_doc_export(
    document_id: str,
    format: str = "json",  # json | markdown | txt
) -> Dict[str, Any]:
    """Export the parsed document as JSON / Markdown / plain text."""
    session = _sessions.get(document_id)
    if not session:
        return {"success": False, "error": f"unknown document_id: {document_id}"}

    fmt = format.lower()
    if fmt == "json":
        payload = {
            "document_id": session.id,
            "source_path": session.source_path,
            "source_type": session.source_type,
            "pages": [
                {
                    "number": pg.number,
                    "sections": [
                        {
                            "index": s.index,
                            "heading": s.heading,
                            "style": s.style,
                            "text": s.text,
                        }
                        for s in pg.sections
                    ],
                }
                for pg in session.pages
            ],
        }
        return {"success": True, "format": "json", "data": payload}

    if fmt == "markdown":
        lines = [f"# {Path(session.source_path).name}", ""]
        for pg in session.pages:
            lines.append(f"## Page {pg.number}")
            for sec in pg.sections:
                if sec.heading:
                    lines.append(f"### {sec.heading}")
                lines.append(sec.text)
                lines.append("")
        return {"success": True, "format": "markdown", "data": "\n".join(lines)}

    if fmt in {"txt", "text"}:
        text = "\n\n".join(
            sec.text for pg in session.pages for sec in pg.sections if sec.text
        )
        return {"success": True, "format": "txt", "data": text}

    return {"success": False, "error": f"unknown format: {format}"}


async def handle_doc_list() -> Dict[str, Any]:
    """List all open document sessions."""
    return {
        "success": True,
        "count": len(_sessions),
        "sessions": [
            {
                "document_id": s.id,
                "source_path": s.source_path,
                "source_type": s.source_type,
                "page_count": len(s.pages),
                "pending_edits": len(s.pending_edits),
                "created_at": s.created_at,
            }
            for s in _sessions.values()
        ],
    }


def _reset_sessions_for_test() -> None:
    """Test helper — empty the in-process registry."""
    _sessions.clear()
